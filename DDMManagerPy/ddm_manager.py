"""
DDM Manager — Metalfrio Solutions
v5.0 CustomTkinter — UI moderna, performance otimizada
"""

# ─── DPI awareness (antes de qualquer import de UI) ──────────────────────────
import ctypes, sys
try:
    ctypes.windll.shcore.SetProcessDpiAwareness(2)
except Exception:
    try:
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass

# ─── Imports padrão (leves, carregam rápido) ─────────────────────────────────
import os, copy, json, hashlib, datetime, threading, subprocess
from pathlib import Path

# ─── UI (CustomTkinter) ───────────────────────────────────────────────────────
import customtkinter as ctk
from tkinter import filedialog, messagebox

# ─── Constantes de tema ───────────────────────────────────────────────────────
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Paleta Metalfrio
BG_APP    = "#0A0F1E"
BG_CARD   = "#0F1729"
BG_INPUT  = "#151E35"
BG_HDR    = "#060D1A"
ACCENT    = "#00D4FF"
ACCENT2   = "#0EA5E9"
SUCCESS   = "#10B981"
WARN      = "#F59E0B"
ERROR     = "#EF4444"
GOLD      = "#FBBF24"
TXT_PRI   = "#F1F5F9"
TXT_SEC   = "#94A3B8"
TXT_DIM   = "#475569"
BORDER    = "#1E293B"

FONT_TITLE  = ("Segoe UI", 20, "bold")
FONT_SUB    = ("Segoe UI", 10)
FONT_LABEL  = ("Segoe UI", 9,  "bold")
FONT_FIELD  = ("Segoe UI", 11)
FONT_BTN    = ("Segoe UI", 12, "bold")
FONT_BTN_SM = ("Segoe UI", 10, "bold")
FONT_LOG    = ("Cascadia Code", 9) if sys.platform == "win32" else ("Consolas", 9)
FONT_SMALL  = ("Segoe UI", 9)

# ─── Mapeamentos ──────────────────────────────────────────────────────────────
DIA_MAP  = {0:"2", 1:"3", 2:"4", 3:"5", 4:"6", 5:"7", 6:"7"}
DIAS_PT  = ["Segunda","Terça","Quarta","Quinta","Sexta","Sábado","Domingo"]
DIAS_ABR = ["SEG","TER","QUA","QUI","SEX","SÁB","DOM"]
MES_MAP  = {
    1:"1 - JANEIRO",  2:"2 - FEVEREIRO", 3:"3 - MARÇO",
    4:"4 - ABRIL",    5:"5 - MAIO",      6:"6 - JUNHO",
    7:"7 - JULHO",    8:"8 - AGOSTO",    9:"9 - SETEMBRO",
    10:"10 - OUTUBRO",11:"11 - NOVEMBRO",12:"12 - DEZEMBRO",
}

PARTICIPANTES = [
    ("10347","ADRIANO GARCIA"),
    ("63363","AGATHA LAÍS SALGUEIRO MAROTTI"),
    ("63051","AMANDA FLORÊNCIO"),
    ("63175","ANA EDUARDA SANO SILVA"),
    ("10934","ANDRÉ LUIZ HIGA FREITAS"),
    ("63538","AUGUSTO GRAÇA SILVESTRIN"),
    ("63231","BARBARA RAMIRES IANHES"),
    ("62334","CLEYTON FARIA ASSIS"),
    ("62397","JOÃO VITOR RODRIGUES SOUZA"),
    ("184",  "JOEL BORGES DE CAMPOS JUNIOR"),
    ("61510","KAROLINE PEDROSO SANTOS"),
    ("10437","LETICIA FERNANDES GONÇALVES"),
    ("63186","LUDIMILA SANTOS SOUZA"),
    ("63516","LUIS GUILHERME DUENHAS SILVA"),
    ("63391","RENAN DA SILVA MANTOVANI"),
    ("63999","RODRIGO ARIAS SILVA"),
    ("11461","VANESSA DIAS DA SILVA"),
    ("9018", "WAGNER JUNIOR ALMEIDA"),
]

DEFAULTS = {
    "raiz":       r"Q:\Transferencia\DDM 2026",
    "setor":      "Engenharia Industrial",
    "subarea":    "Engenharia de Produto",
    "turno":      "ADM",
    "facilitador":"Leitura Individual",
}

# ─── Caminhos de dados ────────────────────────────────────────────────────────
BASE_DIR      = Path(sys.executable if getattr(sys,"frozen",False) else __file__).parent
SETTINGS_FILE = BASE_DIR / "appsettings.json"
LOG_FILE      = BASE_DIR / "ddm_historico.csv"

# ════════════════════════════════════════════════════════════════════════════
# Serviços (sem UI — importados lazy quando necessário)
# ════════════════════════════════════════════════════════════════════════════

def semana_atual(data: datetime.date) -> tuple:
    diff = (data.weekday()) % 7
    seg  = data - datetime.timedelta(days=diff)
    return seg, seg + datetime.timedelta(days=4)

def varrer_semana(raiz: str, data: datetime.date) -> list[dict]:
    resultado = []
    p = Path(raiz)
    if not p.exists():
        return resultado
    seg, sex = semana_atual(data)
    hoje_num = DIA_MAP.get(data.weekday(), "")

    for pasta in sorted(p.iterdir()):
        if not pasta.is_dir(): continue
        partes = pasta.name.split(" - ", 1)
        if len(partes) < 2: continue
        num = partes[0].strip()
        if num not in {"2","3","4","5","6","7"}: continue

        weekday_idx = {"2":0,"3":1,"4":2,"5":3,"6":4,"7":5}[num]
        dia_nome = DIAS_PT[weekday_idx]
        dia_abr  = DIAS_ABR[weekday_idx]

        nome_mes = MES_MAP[data.month]
        pasta_mes = next(
            (m for m in pasta.iterdir()
             if m.is_dir() and (m.name.upper() == nome_mes.upper()
                                or m.name.startswith(f"{data.month} -"))),
            None
        )

        docx_path, tema, data_ddm, erro = None, "—", "", ""
        if not pasta_mes:
            erro = f"Pasta '{nome_mes}' não encontrada"
        else:
            docxs = sorted(pasta_mes.glob("*.docx"))
            docxs = [f for f in docxs if "_PREENCHIDO" not in f.name]
            alvo, melhor = None, None
            for f in docxs:
                try:
                    dd, mm = int(f.name[:2]), int(f.name[3:5])
                    d_arq  = datetime.date(data.year, mm, dd)
                    if seg <= d_arq <= sex:
                        delta = abs((d_arq - data).days)
                        if melhor is None or delta < melhor:
                            alvo, melhor = f, delta
                except: continue
            if alvo:
                docx_path = str(alvo)
                tema      = _extrair_tema(alvo.name)
                data_ddm  = _extrair_data(alvo.name, data.year)
            else:
                erro = f"Sem DDM {seg:%d/%m}–{sex:%d/%m}"

        resultado.append({
            "num":      int(num),
            "dia_nome": dia_nome,
            "dia_abr":  dia_abr,
            "docx":     docx_path,
            "tema":     tema,
            "data_ddm": data_ddm,
            "erro":     erro,
            "eh_hoje":  num == hoje_num,
            "selecionado": docx_path is not None,
        })
    return sorted(resultado, key=lambda x: x["num"])

def _extrair_tema(nome: str) -> str:
    stem = Path(nome).stem
    p = stem.split(" ", 1)
    return p[1].strip() if len(p) == 2 else stem

def _extrair_data(nome: str, ano: int) -> str:
    try:
        dd, mm = int(nome[:2]), int(nome[3:5])
        return f"{dd:02d}/{mm:02d}/{ano}"
    except: return ""

def processar_ddm(src: str, dst: str, campos: dict) -> int:
    """Preenche cabeçalho e participantes. Retorna nomes injetados."""
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    LABELS = {
        "SETOR/LINHA:": campos["setor"],
        "TURNO:":       campos["turno"],
        "SUBÁREA:":     campos["subarea"],
        "FACILITADOR:": campos["facilitador"],
    }

    doc = Document(src)

    # Verifica DATA
    data_no_doc = any(
        c.text.strip().startswith("DATA:") and c.text.strip()[5:].strip()
        for t in doc.tables for row in t.rows
        for cu in _cells_unicas(row) for c in cu.paragraphs
    )
    if not data_no_doc:
        LABELS["DATA:"] = campos["data"]

    tabela  = doc.tables[0]
    part_i  = 0
    nomes   = 0

    for linha in tabela.rows:
        cells = _cells_unicas(linha)

        # Cabeçalho
        for cell in cells:
            for para in cell.paragraphs:
                txt = para.text.strip()
                for label, valor in LABELS.items():
                    if txt.startswith(label) and txt[len(label):].strip() == "":
                        runs = para.runs
                        if not runs: break
                        for r in runs[1:]: r._element.getparent().remove(r._element)
                        novo = copy.deepcopy(runs[0]._element)
                        t = novo.find(qn("w:t"))
                        if t is None: t = etree.SubElement(novo, qn("w:t"))
                        t.text = " " + valor
                        t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
                        para._element.append(novo)
                        break

        # Participante
        if not cells: continue
        try: int(cells[0].text.strip())
        except ValueError: continue

        if part_i < len(PARTICIPANTES) and len(cells) >= 3:
            re_v, nome_v = PARTICIPANTES[part_i]
            _escrever_celula(cells[1], re_v,   "center")
            _escrever_celula(cells[2], nome_v, "left")
            nomes += 1
        part_i += 1

    doc.save(dst)
    return nomes

def _cells_unicas(linha):
    seen, res = set(), []
    for c in linha.cells:
        cid = id(c._tc)
        if cid not in seen:
            seen.add(cid); res.append(c)
    return res

def _escrever_celula(cell, valor: str, alinhar: str):
    from docx.oxml.ns import qn
    from lxml import etree
    tcp = cell._tc.find(qn("w:tcPr"))
    if tcp is None: tcp = etree.SubElement(cell._tc, qn("w:tcPr"))
    va = tcp.find(qn("w:vAlign"))
    if va is None: va = etree.SubElement(tcp, qn("w:vAlign"))
    va.set(qn("w:val"), "center")
    para = cell.paragraphs[0] if cell.paragraphs else None
    if not para: return
    for r in para.runs: r._element.getparent().remove(r._element)
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(para._element, qn("w:pPr"))
        para._element.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None: jc = etree.SubElement(ppr, qn("w:jc"))
    jc.set(qn("w:val"), "center" if alinhar == "center" else "left")
    r_el = etree.SubElement(para._element, qn("w:r"))
    rpr  = etree.SubElement(r_el, qn("w:rPr"))
    fnts = etree.SubElement(rpr, qn("w:rFonts"))
    for a,v in [("w:ascii","Arial"),("w:hAnsi","Arial"),
                ("w:eastAsia","Batang"),("w:cs","Arial")]:
        fnts.set(qn(a), v)
    for tag in ("w:sz","w:szCs"):
        e = etree.SubElement(rpr, qn(tag)); e.set(qn("w:val"), "18")
    t = etree.SubElement(r_el, qn("w:t"))
    t.text = valor
    t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")

def converter_pdfs_word(caminhos: list) -> list:
    try: import win32com.client
    except ImportError: raise RuntimeError("pywin32 não encontrado.")
    import win32com.client as wc
    from pathlib import Path as P
    from ctypes import windll
    resultados = []
    word = None
    try:
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        for c in caminhos:
            c   = P(c).resolve()
            pdf = c.with_suffix(".pdf")
            doc = None
            try:
                doc = word.Documents.Open(str(c), ReadOnly=True, AddToRecentFiles=False)
                doc.SaveAs(str(pdf), FileFormat=17)
                resultados.append((str(pdf), ""))
            except Exception as e:
                resultados.append((None, str(e)))
            finally:
                if doc:
                    try: doc.Close(False)
                    except: pass
    finally:
        if word:
            try: word.Quit()
            except: pass
    return resultados

def combinar_docx_word(docxs: list, pasta: str) -> str:
    try: import win32com.client as wc
    except ImportError: raise RuntimeError("pywin32 não encontrado.")
    from pathlib import Path as P
    saida = str(P(pasta) / "DDM_SEMANA_COMBINADO.docx")
    word  = None
    doc   = None
    try:
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(str(P(docxs[0]).resolve()),
                                   ReadOnly=False, AddToRecentFiles=False)
        for path in docxs[1:]:
            rng = doc.Range(); rng.Collapse(0)
            rng.InsertBreak(7)
            rng = doc.Range(); rng.Collapse(0)
            rng.InsertFile(str(P(path).resolve()))
        doc.SaveAs(saida)
    finally:
        if doc:
            try: doc.Close(False)
            except: pass
        if word:
            try: word.Quit()
            except: pass
    return saida

def imprimir_docxs(caminhos: list, impressora: str | None) -> list:
    try: import win32com.client as wc
    except ImportError: raise RuntimeError("pywin32 não encontrado.")
    resultados = []
    word = None
    try:
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        orig = word.ActivePrinter
        if impressora: word.ActivePrinter = impressora
        for path in caminhos:
            doc = None
            try:
                doc = word.Documents.Open(path, ReadOnly=True, AddToRecentFiles=False)
                doc.PrintOut(Background=False)
                resultados.append((path, True, ""))
            except Exception as e:
                resultados.append((path, False, str(e)))
            finally:
                if doc:
                    try: doc.Close(False)
                    except: pass
        if impressora: word.ActivePrinter = orig
    finally:
        if word:
            try: word.Quit()
            except: pass
    return resultados

def listar_impressoras() -> list[str]:
    try:
        import winreg
        impressoras = []
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
            r"SYSTEM\CurrentControlSet\Control\Print\Printers")
        i = 0
        while True:
            try: impressoras.append(winreg.EnumKey(key, i)); i += 1
            except WindowsError: break
        return sorted(impressoras)
    except: return []

def gerar_xlsx(itens: list, pasta: str, data_ref: datetime.date) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    seg, sex = semana_atual(data_ref)
    wb = Workbook()
    ws = wb.active
    ws.title = "DDMs da Semana"
    azul = "003F7F"; branco = "FFFFFF"

    ws.column_dimensions["A"].width = 16
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 16
    ws.column_dimensions["D"].width = 18

    # Título
    ws.merge_cells("A1:D1")
    ws["A1"] = f"Relatório DDM — Semana {seg:%d/%m} a {sex:%d/%m/%Y}"
    ws["A1"].font = Font(bold=True, size=13, color=azul)
    ws["A2"] = f"Gerado em: {datetime.datetime.now():%d/%m/%Y %H:%M}"
    ws["A2"].font = Font(size=10, color="666666")

    # Cabeçalho
    for col, txt in enumerate(["Dia","Arquivo","Data DDM","Participantes"], 1):
        cell = ws.cell(row=4, column=col, value=txt)
        cell.font      = Font(bold=True, color=branco, size=11)
        cell.fill      = PatternFill("solid", fgColor=azul)
        cell.alignment = Alignment(horizontal="center")

    for i, (dia, arq, data, part) in enumerate(itens, 5):
        ws.cell(row=i, column=1, value=dia)
        ws.cell(row=i, column=2, value=Path(arq).name)
        ws.cell(row=i, column=3, value=data)
        ws.cell(row=i, column=4, value=part)

    saida = str(Path(pasta) / f"Relatorio_DDM_{seg:%d-%m-%Y}.xlsx")
    wb.save(saida)
    return saida

def registrar_log(dia: str, arq: str, tipo: str, nomes: int, ok: bool, erro: str = ""):
    try:
        if not LOG_FILE.exists():
            LOG_FILE.write_text("Data;Hora;Dia;Arquivo;Tipo;Participantes;Status;Erro\n",
                                encoding="utf-8")
        agora = datetime.datetime.now()
        linha = ";".join([
            agora.strftime("%d/%m/%Y"), agora.strftime("%H:%M:%S"),
            dia, Path(arq).name, tipo, str(nomes),
            "OK" if ok else "ERRO", erro.replace(";",",").replace("\n"," ")
        ])
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(linha + "\n")
    except: pass

def carregar_settings() -> dict:
    try:
        if SETTINGS_FILE.exists():
            return json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
    except: pass
    return {**DEFAULTS, "pasta_saida":"", "semana_offset":0}

def salvar_settings(s: dict):
    try: SETTINGS_FILE.write_text(json.dumps(s, ensure_ascii=False, indent=2), encoding="utf-8")
    except: pass

def verificar_word() -> bool:
    try:
        import winreg
        k = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
            r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE")
        k.Close(); return True
    except: return False

def ativar_lembrete() -> tuple[bool, str]:
    try:
        script = (BASE_DIR / "lembrete_ddm.ps1")
        script.write_text("""
$xml = [Windows.UI.Notifications.ToastNotificationManager,Windows.UI.Notifications,ContentType=WindowsRuntime]::GetTemplateContent(
    [Windows.UI.Notifications.ToastTemplateType]::ToastText02)
$xml.GetElementsByTagName('text')[0].AppendChild($xml.CreateTextNode('DDM Manager')) | Out-Null
$xml.GetElementsByTagName('text')[1].AppendChild($xml.CreateTextNode('Gere os DDMs da semana!')) | Out-Null
[Windows.UI.Notifications.ToastNotificationManager]::CreateToastNotifier('DDMManager').Show(
    [Windows.UI.Notifications.ToastNotification]::new($xml))
""", encoding="utf-8")
        r = subprocess.run(
            ["schtasks","/Create","/F","/TN","DDMManager_Lembrete",
             "/TR",f'powershell.exe -WindowStyle Hidden -ExecutionPolicy Bypass -File "{script}"',
             "/SC","WEEKLY","/D","MON","/ST","08:00"],
            capture_output=True, text=True)
        return r.returncode == 0, r.stderr
    except Exception as e:
        return False, str(e)

def desativar_lembrete() -> tuple[bool, str]:
    try:
        r = subprocess.run(
            ["schtasks","/Delete","/F","/TN","DDMManager_Lembrete"],
            capture_output=True, text=True)
        return r.returncode == 0, r.stderr
    except Exception as e:
        return False, str(e)

def lembrete_ativo() -> bool:
    try:
        r = subprocess.run(
            ["schtasks","/Query","/TN","DDMManager_Lembrete"],
            capture_output=True, text=True)
        return r.returncode == 0
    except: return False

# ════════════════════════════════════════════════════════════════════════════
# Interface — CustomTkinter
# ════════════════════════════════════════════════════════════════════════════

class DDMApp(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.title("DDM Manager — Metalfrio")
        self.configure(fg_color=BG_APP)
        self.minsize(1060, 680)
        self._center(1160, 740)

        # Estado
        self._settings      = carregar_settings()
        self._semana_offset = self._settings.get("semana_offset", 0)
        self._pasta_saida   = self._settings.get("pasta_saida", "")
        self._itens: list[dict] = []
        self._checks: list[ctk.BooleanVar] = []
        self._check_widgets: list = []
        self._em_operacao   = False

        self._build_ui()
        self._carregar_preferencias()

        # Carrega dados em background após UI estar pronta
        self.after(50, self._varrer_async)
        self.after(100, self._carregar_impressoras_async)
        self.after(150, self._iniciar_relogio)
        self.protocol("WM_DELETE_WINDOW", self._sair)

    def _center(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth()  - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    # ── BUILD UI ─────────────────────────────────────────────────────────────

    def _build_ui(self):
        self.grid_rowconfigure(2, weight=1)
        self.grid_columnconfigure(0, weight=1)

        self._build_header()
        self._build_paths_bar()
        self._build_body()
        self._build_footer()

    # ── Header ───────────────────────────────────────────────────────────────

    def _build_header(self):
        hdr = ctk.CTkFrame(self, fg_color=BG_HDR, corner_radius=0, height=64)
        hdr.grid(row=0, column=0, sticky="ew")
        hdr.grid_propagate(False)
        hdr.grid_columnconfigure(1, weight=1)

        # Logo
        logo = ctk.CTkFrame(hdr, fg_color="transparent")
        logo.grid(row=0, column=0, padx=20, pady=10, sticky="w")
        ctk.CTkLabel(logo, text="❄", font=("Segoe UI", 28),
                     text_color=ACCENT).pack(side="left", padx=(0,10))
        titles = ctk.CTkFrame(logo, fg_color="transparent")
        titles.pack(side="left")
        ctk.CTkLabel(titles, text="DDM Manager",
                     font=FONT_TITLE, text_color=TXT_PRI).pack(anchor="w")
        ctk.CTkLabel(titles, text="Metalfrio Solutions  ·  Engenharia de Produto",
                     font=FONT_SMALL, text_color=TXT_SEC).pack(anchor="w")

        # Dica atalhos + relógio
        right = ctk.CTkFrame(hdr, fg_color="transparent")
        right.grid(row=0, column=2, padx=16, pady=10, sticky="e")
        ctk.CTkLabel(right,
                     text="F5 Atualizar  |  Ctrl+P PDF  |  Ctrl+D DOCX  |  Alt+◄► Semana",
                     font=("Segoe UI", 8), text_color=TXT_DIM).pack(anchor="e")
        self._lbl_clock = ctk.CTkLabel(right, text="",
                                       font=FONT_SMALL, text_color=TXT_SEC)
        self._lbl_clock.pack(anchor="e")

        # Accent line
        ctk.CTkFrame(self, fg_color=ACCENT, height=3, corner_radius=0
                     ).grid(row=0, column=0, sticky="sew")

        # Atalhos de teclado
        self.bind("<F5>",         lambda _: self._varrer_async())
        self.bind("<Control-p>",  lambda _: self._gerar_pdf())
        self.bind("<Control-d>",  lambda _: self._gerar_docx())
        self.bind("<Alt-Left>",   lambda _: self._semana_anterior())
        self.bind("<Alt-Right>",  lambda _: self._semana_proxima())

    # ── Paths bar ────────────────────────────────────────────────────────────

    def _build_paths_bar(self):
        bar = ctk.CTkFrame(self, fg_color=BG_CARD, corner_radius=0)
        bar.grid(row=1, column=0, sticky="ew")
        bar.grid_columnconfigure(1, weight=1)

        # Pasta raiz
        raiz_row = ctk.CTkFrame(bar, fg_color="transparent")
        raiz_row.grid(row=0, column=0, columnspan=4, sticky="ew", padx=16, pady=(10,4))
        raiz_row.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(raiz_row, text="📁  PASTA RAIZ DDM 2026",
                     font=FONT_LABEL, text_color=TXT_SEC
                     ).grid(row=0, column=0, padx=(0,12), sticky="w")
        self._v_raiz = ctk.StringVar()
        ctk.CTkEntry(raiz_row, textvariable=self._v_raiz,
                     font=FONT_FIELD, fg_color=BG_INPUT,
                     border_color=BORDER, text_color=TXT_PRI
                     ).grid(row=0, column=1, sticky="ew", ipady=2)
        ctk.CTkButton(raiz_row, text="  …  ", width=50,
                      font=FONT_BTN_SM, fg_color=ACCENT, text_color=BG_HDR,
                      hover_color=ACCENT2, command=self._browse_raiz
                      ).grid(row=0, column=2, padx=(6,4))
        ctk.CTkButton(raiz_row, text="🔄  Atualizar  [F5]", width=160,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER, border_width=1, border_color=ACCENT,
                      command=self._varrer_async
                      ).grid(row=0, column=3)

        # Pasta saída
        saida_row = ctk.CTkFrame(bar, fg_color="transparent")
        saida_row.grid(row=1, column=0, columnspan=4, sticky="ew", padx=16, pady=(0,10))
        saida_row.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(saida_row, text="💾  SALVAR EM",
                     font=FONT_LABEL, text_color=TXT_SEC
                     ).grid(row=0, column=0, padx=(0,12), sticky="w")
        self._lbl_saida = ctk.CTkLabel(saida_row,
                     text="(padrão: mesma pasta do DDM)",
                     font=FONT_SMALL, text_color=TXT_DIM, anchor="w")
        self._lbl_saida.grid(row=0, column=1, sticky="ew")
        ctk.CTkButton(saida_row, text="Alterar", width=80,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER, command=self._browse_saida
                      ).grid(row=0, column=2, padx=(6,4))
        ctk.CTkButton(saida_row, text="✕", width=40,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ERROR,
                      hover_color=BORDER, command=self._limpar_saida
                      ).grid(row=0, column=3)

    # ── Body ─────────────────────────────────────────────────────────────────

    def _build_body(self):
        body = ctk.CTkFrame(self, fg_color=BG_APP, corner_radius=0)
        body.grid(row=2, column=0, sticky="nsew")
        body.grid_columnconfigure(1, weight=0)
        body.grid_columnconfigure(2, weight=0)
        body.grid_columnconfigure(3, weight=0)
        body.grid_columnconfigure(4, weight=1)
        body.grid_rowconfigure(0, weight=1)

        self._build_col_campos(body)
        _sep(body, col=1)
        self._build_col_lista(body)
        _sep(body, col=3)
        self._build_col_preview_log(body)

    def _build_col_campos(self, parent):
        col = ctk.CTkScrollableFrame(parent, fg_color=BG_CARD,
                                     corner_radius=0, width=290)
        col.grid(row=0, column=0, sticky="nsew", padx=(0,0))

        _section(col, "CABEÇALHO DDM")

        self._v_setor       = ctk.StringVar()
        self._v_subarea     = ctk.StringVar()
        self._v_turno       = ctk.StringVar()
        self._v_facilitador = ctk.StringVar()

        for label, var, icon in [
            ("SETOR / LINHA",  self._v_setor,       "🏭"),
            ("SUBÁREA",        self._v_subarea,     "🔹"),
            ("TURNO",          self._v_turno,       "⏰"),
            ("FACILITADOR",    self._v_facilitador, "👤"),
        ]:
            ctk.CTkLabel(col, text=f"{icon}  {label}",
                         font=FONT_LABEL, text_color=TXT_SEC, anchor="w"
                         ).pack(fill="x", padx=14, pady=(10,3))
            e = ctk.CTkEntry(col, textvariable=var,
                             font=FONT_FIELD, fg_color=BG_INPUT,
                             border_color=BORDER, text_color=TXT_PRI)
            e.pack(fill="x", padx=14, ipady=4)
            var.trace_add("write", lambda *_: self._atualizar_checklist())

        # Badge participantes
        badge = ctk.CTkFrame(col, fg_color=BG_INPUT, corner_radius=8)
        badge.pack(fill="x", padx=14, pady=(16,0))
        ctk.CTkLabel(badge,
                     text=f"👥  {len(PARTICIPANTES)} participantes carregados",
                     font=FONT_SMALL, text_color=SUCCESS
                     ).pack(padx=12, pady=8)

        # Checklist
        _section(col, "VERIFICAÇÕES")
        self._chk_pasta   = _chk_label(col)
        self._chk_word    = _chk_label(col)
        self._chk_ddms    = _chk_label(col)
        self._chk_campos  = _chk_label(col)

        # Impressora
        _section(col, "IMPRESSORA")
        self._v_impressora = ctk.StringVar(value="(padrão do sistema)")
        self._cmb_impressora = ctk.CTkComboBox(col, variable=self._v_impressora,
                                               font=FONT_SMALL,
                                               fg_color=BG_INPUT,
                                               border_color=BORDER,
                                               button_color=ACCENT,
                                               values=["(padrão do sistema)"])
        self._cmb_impressora.pack(fill="x", padx=14, pady=(4,0))

        # Lembrete semanal
        _section(col, "LEMBRETE SEMANAL")
        self._lbl_notif = ctk.CTkLabel(col, text="Verificando...",
                                       font=FONT_SMALL, text_color=TXT_SEC,
                                       anchor="w", wraplength=240)
        self._lbl_notif.pack(fill="x", padx=14, pady=(4,6))
        btn_notif = ctk.CTkFrame(col, fg_color="transparent")
        btn_notif.pack(fill="x", padx=14)
        ctk.CTkButton(btn_notif, text="Ativar", width=90,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER, command=self._ativar_lembrete
                      ).pack(side="left", padx=(0,6))
        ctk.CTkButton(btn_notif, text="Desativar", width=90,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ERROR,
                      hover_color=BORDER, command=self._desativar_lembrete
                      ).pack(side="left")

    def _build_col_lista(self, parent):
        col = ctk.CTkFrame(parent, fg_color=BG_APP, corner_radius=0, width=400)
        col.grid(row=0, column=2, sticky="nsew", padx=0)
        col.grid_rowconfigure(2, weight=1)
        col.grid_columnconfigure(0, weight=1)
        col.grid_propagate(False)

        # Navegação semana
        nav = ctk.CTkFrame(col, fg_color="transparent")
        nav.grid(row=0, column=0, sticky="ew", padx=14, pady=(12,6))
        nav.grid_columnconfigure(1, weight=1)

        ctk.CTkButton(nav, text="◄  [Alt+←]", width=110,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER, command=self._semana_anterior
                      ).grid(row=0, column=0)

        semana_info = ctk.CTkFrame(nav, fg_color="transparent")
        semana_info.grid(row=0, column=1, padx=6)
        self._lbl_semana       = ctk.CTkLabel(semana_info, text="",
                                              font=("Segoe UI",11,"bold"),
                                              text_color=ACCENT)
        self._lbl_semana.pack()
        self._lbl_semana_label = ctk.CTkLabel(semana_info, text="",
                                              font=FONT_SMALL, text_color=TXT_DIM)
        self._lbl_semana_label.pack()

        ctk.CTkButton(nav, text="[Alt+►]  ►", width=110,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER, command=self._semana_proxima
                      ).grid(row=0, column=2)

        # Cabeçalho lista
        hdr = ctk.CTkFrame(col, fg_color=BORDER, corner_radius=4, height=28)
        hdr.grid(row=1, column=0, sticky="ew", padx=14, pady=(0,2))
        hdr.grid_propagate(False)
        for txt, w, anchor in [("✔",3,"center"),("DIA",6,"w"),
                                ("TEMA",20,"w"),("DATA",8,"e")]:
            ctk.CTkLabel(hdr, text=txt, font=FONT_LABEL,
                         text_color=TXT_DIM, width=w, anchor=anchor
                         ).pack(side="left", padx=6)

        # Lista scrollável
        self._lista_frame = ctk.CTkScrollableFrame(
            col, fg_color=BG_APP, corner_radius=0)
        self._lista_frame.grid(row=2, column=0, sticky="nsew", padx=0)

        # Botões seleção
        sel_row = ctk.CTkFrame(col, fg_color="transparent")
        sel_row.grid(row=3, column=0, sticky="ew", padx=14, pady=8)
        ctk.CTkButton(sel_row, text="Selecionar tudo", width=140,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=ACCENT,
                      hover_color=BORDER,
                      command=lambda: self._set_todos(True)
                      ).pack(side="left", padx=(0,6))
        ctk.CTkButton(sel_row, text="Limpar seleção", width=130,
                      font=FONT_BTN_SM, fg_color=BG_INPUT, text_color=TXT_DIM,
                      hover_color=BORDER,
                      command=lambda: self._set_todos(False)
                      ).pack(side="left")

    def _build_col_preview_log(self, parent):
        col = ctk.CTkFrame(parent, fg_color=BG_APP, corner_radius=0)
        col.grid(row=0, column=4, sticky="nsew")
        col.grid_rowconfigure(1, weight=1)
        col.grid_columnconfigure(0, weight=1)

        # Preview
        prev_frame = ctk.CTkFrame(col, fg_color=BG_CARD, corner_radius=8)
        prev_frame.grid(row=0, column=0, sticky="ew", padx=12, pady=(12,6))
        _section(prev_frame, "PREVIEW")
        self._lbl_prev_dia    = ctk.CTkLabel(prev_frame, text="—",
                                             font=("Segoe UI",12,"bold"),
                                             text_color=ACCENT, anchor="w")
        self._lbl_prev_dia.pack(fill="x", padx=14, pady=(0,4))
        self._lbl_prev_tema   = ctk.CTkLabel(prev_frame, text="Selecione um DDM",
                                             font=FONT_SMALL, text_color=TXT_SEC,
                                             anchor="w", wraplength=280)
        self._lbl_prev_tema.pack(fill="x", padx=14)
        self._lbl_prev_data   = ctk.CTkLabel(prev_frame, text="",
                                             font=FONT_SMALL, text_color=TXT_DIM,
                                             anchor="w")
        self._lbl_prev_data.pack(fill="x", padx=14, pady=(2,8))

        # Log
        log_hdr = ctk.CTkFrame(col, fg_color=BG_CARD, corner_radius=0, height=32)
        log_hdr.grid(row=1, column=0, sticky="new")
        log_hdr.grid_propagate(False)
        log_hdr.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(log_hdr, text="LOG DE EXECUÇÃO",
                     font=FONT_LABEL, text_color=TXT_DIM
                     ).grid(row=0, column=0, padx=12, pady=4, sticky="w")
        ctk.CTkButton(log_hdr, text="Limpar", width=70,
                      font=FONT_BTN_SM, fg_color="transparent",
                      text_color=ERROR, hover_color=BG_CARD,
                      command=self._limpar_log
                      ).grid(row=0, column=1, padx=6, pady=2)
        ctk.CTkButton(log_hdr, text="Histórico", width=80,
                      font=FONT_BTN_SM, fg_color="transparent",
                      text_color=ACCENT, hover_color=BG_CARD,
                      command=lambda: os.startfile(str(LOG_FILE))
                                      if LOG_FILE.exists() else None
                      ).grid(row=0, column=2, padx=(0,8), pady=2)

        self._log_text = ctk.CTkTextbox(col, font=FONT_LOG,
                                        fg_color=BG_APP, text_color=TXT_PRI,
                                        corner_radius=0, wrap="word",
                                        activate_scrollbars=True)
        self._log_text.grid(row=1, column=0, sticky="nsew", pady=(32,0))
        self._log_text.configure(state="disabled")

        # Barra de progresso
        self._prog_frame = ctk.CTkFrame(col, fg_color=BG_CARD,
                                        corner_radius=0, height=36)
        self._prog_frame.grid(row=2, column=0, sticky="ew")
        self._prog_frame.grid_propagate(False)
        self._prog_frame.grid_columnconfigure(0, weight=1)
        self._lbl_prog = ctk.CTkLabel(self._prog_frame, text="",
                                      font=FONT_SMALL, text_color=TXT_SEC)
        self._lbl_prog.grid(row=0, column=0, padx=12, pady=2, sticky="w")
        self._progressbar = ctk.CTkProgressBar(self._prog_frame,
                                               mode="indeterminate",
                                               progress_color=ACCENT,
                                               fg_color=BORDER)
        self._progressbar.grid(row=1, column=0, sticky="ew", padx=12, pady=(0,4))
        self._prog_frame.grid_remove()

    # ── Footer ───────────────────────────────────────────────────────────────

    def _build_footer(self):
        footer = ctk.CTkFrame(self, fg_color=BG_HDR, corner_radius=0, height=62)
        footer.grid(row=3, column=0, sticky="ew")
        footer.grid_propagate(False)
        footer.grid_columnconfigure(4, weight=1)

        pad = {"padx": 6, "pady": 12}
        ctk.CTkButton(footer, text="📄   Gerar DOCX  [Ctrl+D]",
                      font=FONT_BTN, fg_color=BG_CARD, text_color=ACCENT,
                      hover_color=BORDER, border_width=1, border_color=ACCENT,
                      command=self._gerar_docx
                      ).grid(row=0, column=0, **pad, padx=(12,6))
        ctk.CTkButton(footer, text="📋   Combinar DOCX",
                      font=FONT_BTN, fg_color=BG_CARD, text_color=ACCENT2,
                      hover_color=BORDER, border_width=1, border_color=ACCENT2,
                      command=self._combinar_docx
                      ).grid(row=0, column=1, **pad)
        ctk.CTkButton(footer, text="📑   Gerar PDF + Abrir  [Ctrl+P]",
                      font=FONT_BTN, fg_color=ACCENT, text_color=BG_HDR,
                      hover_color=ACCENT2,
                      command=self._gerar_pdf
                      ).grid(row=0, column=2, **pad)
        ctk.CTkButton(footer, text="🖨️   Imprimir direto",
                      font=FONT_BTN, fg_color=BG_CARD, text_color=SUCCESS,
                      hover_color=BORDER, border_width=1, border_color=SUCCESS,
                      command=self._imprimir
                      ).grid(row=0, column=3, **pad)
        ctk.CTkButton(footer, text="📊   Relatório",
                      font=FONT_BTN, fg_color=BG_CARD, text_color=GOLD,
                      hover_color=BORDER, border_width=1, border_color=GOLD,
                      command=self._relatorio
                      ).grid(row=0, column=4, **pad, sticky="w")
        ctk.CTkButton(footer, text="✖   Sair",
                      font=FONT_BTN, fg_color=BG_CARD, text_color=ERROR,
                      hover_color=BORDER,
                      command=self._sair
                      ).grid(row=0, column=5, **pad, padx=(6,12), sticky="e")

    # ── Preferências ─────────────────────────────────────────────────────────

    def _carregar_preferencias(self):
        s = self._settings
        self._v_raiz.set(s.get("raiz", DEFAULTS["raiz"]))
        self._v_setor.set(s.get("setor", DEFAULTS["setor"]))
        self._v_subarea.set(s.get("subarea", DEFAULTS["subarea"]))
        self._v_turno.set(s.get("turno", DEFAULTS["turno"]))
        self._v_facilitador.set(s.get("facilitador", DEFAULTS["facilitador"]))
        self._semana_offset = s.get("semana_offset", 0)
        if s.get("pasta_saida"):
            self._pasta_saida = s["pasta_saida"]
            self._lbl_saida.configure(text=s["pasta_saida"], text_color=TXT_PRI)
        self.after(200, self._atualizar_notif)

    def _salvar_preferencias(self):
        salvar_settings({
            "raiz":          self._v_raiz.get(),
            "setor":         self._v_setor.get(),
            "subarea":       self._v_subarea.get(),
            "turno":         self._v_turno.get(),
            "facilitador":   self._v_facilitador.get(),
            "pasta_saida":   self._pasta_saida,
            "semana_offset": self._semana_offset,
        })

    # ── Relógio ───────────────────────────────────────────────────────────────

    def _iniciar_relogio(self):
        def tick():
            n = datetime.datetime.now()
            self._lbl_clock.configure(
                text=f"{DIAS_PT[n.weekday()]}, {n:%d/%m/%Y  %H:%M:%S}")
            self.after(1000, tick)
        tick()

    # ── Varredura ─────────────────────────────────────────────────────────────

    def _varrer_async(self):
        raiz  = self._v_raiz.get().strip()
        hoje  = datetime.date.today()
        data  = hoje + datetime.timedelta(weeks=self._semana_offset)
        seg, sex = semana_atual(data)

        self._lbl_semana.configure(text=f"{seg:%d/%m} – {sex:%d/%m}")
        self._lbl_semana_label.configure(
            text="(semana atual)" if self._semana_offset == 0
            else f"({abs(self._semana_offset)} sem. {'atrás' if self._semana_offset < 0 else 'à frente'})")

        self._limpar_log()
        self._log(f"Semana: {seg:%d/%m/%Y} – {sex:%d/%m/%Y}", TXT_SEC)
        self._log(f"Raiz:   {raiz}\n", TXT_SEC)

        def worker():
            lista = varrer_semana(raiz, data)
            self.after(0, lambda: self._atualizar_lista(lista))

        threading.Thread(target=worker, daemon=True).start()

    def _atualizar_lista(self, lista: list):
        self._itens = lista
        # Limpa widgets anteriores
        for w in self._lista_frame.winfo_children():
            w.destroy()
        self._checks.clear()
        self._check_widgets.clear()

        hoje_num = DIA_MAP.get(datetime.date.today().weekday(), "")

        for item in lista:
            var = ctk.BooleanVar(value=item["selecionado"])
            self._checks.append(var)

            tem   = item["docx"] is not None
            hoje  = item["eh_hoje"]
            bg    = "#0C2545" if hoje else BG_APP
            fg_dia = ACCENT if hoje else (TXT_PRI if tem else TXT_DIM)

            row = ctk.CTkFrame(self._lista_frame, fg_color=bg,
                               corner_radius=4, height=36)
            row.pack(fill="x", pady=1, padx=2)
            row.pack_propagate(False)
            row.grid_columnconfigure(2, weight=1)

            # Borda esquerda para o dia de hoje
            if hoje:
                ctk.CTkFrame(row, fg_color=ACCENT, width=3, corner_radius=0
                             ).grid(row=0, column=0, sticky="ns")

            cb = ctk.CTkCheckBox(row, variable=var, text="",
                                 width=24, checkbox_width=16, checkbox_height=16,
                                 fg_color=ACCENT, hover_color=ACCENT2,
                                 state="normal" if tem else "disabled",
                                 command=lambda i=item, v=var: self._on_check(i, v))
            cb.grid(row=0, column=1, padx=(6,4))

            ctk.CTkLabel(row, text=item["dia_abr"], font=("Segoe UI",10,"bold"),
                         text_color=fg_dia, width=40, anchor="w"
                         ).grid(row=0, column=2, sticky="w")

            tema_txt = (item["tema"][:26]+"…" if len(item["tema"])>26
                        else item["tema"]) if tem else "— sem arquivo —"
            ctk.CTkLabel(row,
                         text=tema_txt, font=FONT_SMALL,
                         text_color=TXT_PRI if tem else TXT_DIM, anchor="w"
                         ).grid(row=0, column=3, sticky="ew", padx=4)

            if item.get("data_ddm"):
                ctk.CTkLabel(row, text=item["data_ddm"], font=FONT_SMALL,
                             text_color=TXT_DIM, width=68, anchor="e"
                             ).grid(row=0, column=4, padx=(0,8))

            # Clique na linha inteira
            if tem:
                for widget in (row,):
                    widget.bind("<Button-1>",
                        lambda e, v=var, i=item: (
                            v.set(not v.get()), self._on_check(i, v)))

        ok  = sum(1 for d in lista if d["docx"])
        err = sum(1 for d in lista if not d["docx"])
        self._log(f"✔  {ok} DDM(s) encontrado(s)", SUCCESS)
        for d in lista:
            if d["docx"]:
                self._log(f"   {d['dia_nome']}: {Path(d['docx']).name}", TXT_SEC)
        if err:
            self._log(f"\n⚠  {err} pasta(s) sem DDM:", WARN)
            for d in lista:
                if not d["docx"]:
                    self._log(f"   {d['dia_nome']}: {d['erro']}", TXT_DIM)

        self._atualizar_checklist()

    def _on_check(self, item: dict, var: ctk.BooleanVar):
        item["selecionado"] = var.get()
        if var.get():
            self._atualizar_preview(item)

    # ── Preview ───────────────────────────────────────────────────────────────

    def _atualizar_preview(self, item: dict):
        self._lbl_prev_dia.configure(text=item["dia_nome"])
        self._lbl_prev_tema.configure(text=item["tema"])
        self._lbl_prev_data.configure(
            text=f"Data: {item['data_ddm']}  |  Facilitador: {self._v_facilitador.get()}")

    # ── Checklist ─────────────────────────────────────────────────────────────

    def _atualizar_checklist(self):
        raiz_ok = Path(self._v_raiz.get()).exists()
        _set_chk(self._chk_pasta, raiz_ok, "Pasta raiz acessível", "Pasta raiz inacessível")

        word_ok = verificar_word()
        _set_chk(self._chk_word, word_ok, "Word instalado", "Word não encontrado")

        n_ddms  = sum(1 for d in self._itens if d["docx"])
        _set_chk(self._chk_ddms, n_ddms > 0, f"{n_ddms} DDM(s) disponível(is)", "Sem DDMs")

        campos_ok = bool(self._v_setor.get() and self._v_facilitador.get())
        _set_chk(self._chk_campos, campos_ok, "Campos preenchidos", "Campos incompletos")

    # ── Impressoras ───────────────────────────────────────────────────────────

    def _carregar_impressoras_async(self):
        def worker():
            impressoras = listar_impressoras()
            self.after(0, lambda: self._cmb_impressora.configure(
                values=["(padrão do sistema)"] + impressoras))
        threading.Thread(target=worker, daemon=True).start()

    # ── Notificação ───────────────────────────────────────────────────────────

    def _atualizar_notif(self):
        ativo = lembrete_ativo()
        self._lbl_notif.configure(
            text="✔  Toda segunda às 8h" if ativo else "Sem lembrete configurado",
            text_color=SUCCESS if ativo else TXT_SEC)

    def _ativar_lembrete(self):
        ok, erro = ativar_lembrete()
        if ok: self._log("✔  Lembrete ativado", SUCCESS)
        else:  self._log(f"✖  {erro}", ERROR)
        self._atualizar_notif()

    def _desativar_lembrete(self):
        ok, erro = desativar_lembrete()
        if ok: self._log("✔  Lembrete desativado", TXT_SEC)
        else:  self._log(f"✖  {erro}", ERROR)
        self._atualizar_notif()

    # ── Helpers UI ────────────────────────────────────────────────────────────

    def _set_todos(self, val: bool):
        for item, var in zip(self._itens, self._checks):
            if item["docx"] or not val:
                var.set(val)
                item["selecionado"] = val

    def _selecionados(self) -> list[dict]:
        sel = [d for d, v in zip(self._itens, self._checks)
               if v.get() and d["docx"]]
        if not sel:
            messagebox.showwarning("Nenhum selecionado",
                "Selecione ao menos um DDM da lista.")
        return sel

    def _pasta_efetiva(self, fallback: str) -> str:
        if self._pasta_saida:
            Path(self._pasta_saida).mkdir(parents=True, exist_ok=True)
            return self._pasta_saida
        return fallback

    def _campos_para(self, item: dict) -> dict:
        return {
            "setor":       self._v_setor.get(),
            "subarea":     self._v_subarea.get(),
            "turno":       self._v_turno.get(),
            "facilitador": self._v_facilitador.get(),
            "data":        item.get("data_ddm") or datetime.date.today().strftime("%d/%m/%Y"),
        }

    def _iniciar_op(self, msg: str):
        self._em_operacao = True
        self._lbl_prog.configure(text=msg)
        self._prog_frame.grid()
        self._progressbar.start()

    def _finalizar_op(self):
        self._em_operacao = False
        self._progressbar.stop()
        self._prog_frame.grid_remove()

    def _gerar_docxs_base(self, sel: list) -> list[str]:
        """Gera DOCXs preenchidos para os itens selecionados. Síncrono (roda em thread)."""
        gerados = []
        for d in sel:
            pasta = self._pasta_efetiva(str(Path(d["docx"]).parent))
            dst   = str(Path(pasta) / (Path(d["docx"]).stem + "_PREENCHIDO.docx"))
            try:
                nomes = processar_ddm(d["docx"], dst, self._campos_para(d))
                self.after(0, lambda dn=d["dia_nome"], f=dst, n=nomes:
                    self._log(f"✔  {dn}: {Path(f).name} ({n} nomes)", SUCCESS))
                registrar_log(d["dia_nome"], dst, "DOCX", nomes, True)
                gerados.append(dst)
            except Exception as e:
                self.after(0, lambda dn=d["dia_nome"], err=str(e):
                    self._log(f"✖  {dn}: {err}", ERROR))
                registrar_log(d["dia_nome"], d["docx"], "DOCX", 0, False, str(e))
        return gerados

    # ── Ações principais ─────────────────────────────────────────────────────

    def _gerar_docx(self):
        sel = self._selecionados()
        if not sel or self._em_operacao: return
        self._iniciar_op(f"Gerando {len(sel)} DOCX(s)...")
        self._log(f"\n─── Gerando {len(sel)} DOCX(s)  [Ctrl+D] ───", ACCENT)

        def worker():
            gerados = self._gerar_docxs_base(sel)
            self.after(0, lambda: self._finalizar_op())
            if gerados:
                self.after(0, lambda: messagebox.askyesno(
                    "Concluído", f"{len(gerados)} DOCX(s) gerado(s).\nAbrir pasta?") and
                    os.startfile(str(Path(gerados[0]).parent)))

        threading.Thread(target=worker, daemon=True).start()

    def _gerar_pdf(self):
        sel = self._selecionados()
        if not sel or self._em_operacao: return
        self._iniciar_op(f"Gerando {len(sel)} PDF(s)...")
        self._log(f"\n─── Gerando {len(sel)} PDF(s)  [Ctrl+P] ───", ACCENT)

        def worker():
            docxs = self._gerar_docxs_base(sel)
            if not docxs:
                self.after(0, self._finalizar_op); return

            self.after(0, lambda: self._log("   Convertendo via Word...", TXT_SEC))
            try:
                resultados = converter_pdfs_word(docxs)
            except Exception as e:
                self.after(0, lambda: self._log(f"✖  Word: {e}", ERROR))
                self.after(0, self._finalizar_op); return

            pdfs = []
            for (pdf, erro), d in zip(resultados, sel):
                if pdf:
                    self.after(0, lambda p=pdf, dn=d["dia_nome"]:
                        self._log(f"✔  {dn}: {Path(p).name}", SUCCESS))
                    registrar_log(d["dia_nome"], pdf, "PDF", 0, True)
                    pdfs.append(pdf)
                else:
                    self.after(0, lambda dn=d["dia_nome"], e=erro:
                        self._log(f"✖  {dn}: {e}", ERROR))

            self.after(0, self._finalizar_op)
            if not pdfs: return

            if len(pdfs) == 1:
                self.after(0, lambda: os.startfile(pdfs[0]))
            else:
                self.after(0, lambda: self._log(f"   Mesclando {len(pdfs)} PDFs...", TXT_SEC))
                try:
                    pasta = self._pasta_efetiva(str(Path(pdfs[0]).parent))
                    merged = combinar_docx_word(docxs, pasta)
                    # Converte o combinado para PDF
                    res = converter_pdfs_word([merged])
                    pdf_final = res[0][0] if res and res[0][0] else None
                    if pdf_final:
                        self.after(0, lambda: self._log(
                            f"✔  PDF combinado: {Path(pdf_final).name}", SUCCESS))
                        self.after(0, lambda: os.startfile(pdf_final))
                    else:
                        for p in pdfs:
                            self.after(0, lambda pp=p: os.startfile(pp))
                except Exception as e:
                    self.after(0, lambda: self._log(f"⚠  Mesclagem: {e}", WARN))
                    for p in pdfs:
                        self.after(0, lambda pp=p: os.startfile(pp))

        threading.Thread(target=worker, daemon=True).start()

    def _combinar_docx(self):
        sel = self._selecionados()
        if not sel or self._em_operacao: return
        if len(sel) < 2:
            messagebox.showinfo("Atenção", "Selecione ao menos 2 DDMs para combinar.")
            return
        self._iniciar_op(f"Combinando {len(sel)} DOCX(s)...")
        self._log(f"\n─── Combinando {len(sel)} DOCX(s) ───", ACCENT)

        def worker():
            docxs = self._gerar_docxs_base(sel)
            if len(docxs) < 2:
                self.after(0, self._finalizar_op); return
            try:
                pasta   = self._pasta_efetiva(str(Path(docxs[0]).parent))
                saida   = combinar_docx_word(docxs, pasta)
                self.after(0, lambda: self._log(
                    f"✔  Combinado: {Path(saida).name}", SUCCESS))
                registrar_log("SEMANA", saida, "DOCX_COMBINADO", 0, True)
                self.after(0, lambda: messagebox.askyesno(
                    "Concluído", f"DOCX combinado:\n{Path(saida).name}\n\nAbrir?") and
                    os.startfile(saida))
            except Exception as e:
                self.after(0, lambda: self._log(f"✖  {e}", ERROR))
            self.after(0, self._finalizar_op)

        threading.Thread(target=worker, daemon=True).start()

    def _imprimir(self):
        sel = self._selecionados()
        if not sel or self._em_operacao: return
        imp = self._v_impressora.get()
        imp = None if imp == "(padrão do sistema)" else imp
        if not messagebox.askyesno("Confirmar",
            f"Imprimir {len(sel)} DDM(s)?\n"
            f"Impressora: {imp or 'padrão do sistema'}"):
            return
        self._iniciar_op("Imprimindo...")
        self._log(f"\n─── Impressão direta: {len(sel)} DDM(s) ───", ACCENT)

        def worker():
            docxs = self._gerar_docxs_base(sel)
            if not docxs:
                self.after(0, self._finalizar_op); return
            try:
                resultados = imprimir_docxs(docxs, imp)
                for arq, ok, erro in resultados:
                    if ok:
                        self.after(0, lambda f=arq:
                            self._log(f"✔  {Path(f).name}: enviado", SUCCESS))
                        registrar_log("IMPRESSÃO", arq, "PRINT", 0, True)
                    else:
                        self.after(0, lambda f=arq, e=erro:
                            self._log(f"✖  {Path(f).name}: {e}", ERROR))
            except Exception as e:
                self.after(0, lambda: self._log(f"✖  {e}", ERROR))
            self.after(0, self._finalizar_op)

        threading.Thread(target=worker, daemon=True).start()

    def _relatorio(self):
        if self._em_operacao: return
        ddms_ok = [d for d in self._itens if d["docx"]]
        if not ddms_ok:
            messagebox.showinfo("Sem dados", "Nenhum DDM disponível.")
            return
        self._iniciar_op("Gerando relatório...")
        self._log("\n─── Relatório semanal ───", ACCENT)

        def worker():
            try:
                hoje    = datetime.date.today()
                data    = hoje + datetime.timedelta(weeks=self._semana_offset)
                pasta   = self._pasta_saida or str(BASE_DIR)
                itens_r = [(d["dia_nome"], d["docx"], d["data_ddm"],
                            len(PARTICIPANTES)) for d in ddms_ok]
                saida   = gerar_xlsx(itens_r, pasta, data)
                self.after(0, lambda: self._log(
                    f"✔  Relatório: {Path(saida).name}", SUCCESS))
                self.after(0, lambda: os.startfile(saida))
            except Exception as e:
                self.after(0, lambda: self._log(f"✖  {e}", ERROR))
            self.after(0, self._finalizar_op)

        threading.Thread(target=worker, daemon=True).start()

    # ── Navegação ─────────────────────────────────────────────────────────────

    def _semana_anterior(self):
        self._semana_offset -= 1; self._varrer_async()

    def _semana_proxima(self):
        self._semana_offset += 1; self._varrer_async()

    # ── Browser de pasta ──────────────────────────────────────────────────────

    def _browse_raiz(self):
        p = filedialog.askdirectory(title="Pasta raiz DDM 2026",
                                    initialdir=self._v_raiz.get())
        if p:
            self._v_raiz.set(p)
            self._varrer_async()

    def _browse_saida(self):
        p = filedialog.askdirectory(title="Onde salvar os arquivos gerados")
        if p:
            self._pasta_saida = p
            self._lbl_saida.configure(text=p, text_color=TXT_PRI)

    def _limpar_saida(self):
        self._pasta_saida = ""
        self._lbl_saida.configure(text="(padrão: mesma pasta do DDM)",
                                  text_color=TXT_DIM)

    # ── Log ───────────────────────────────────────────────────────────────────

    def _log(self, msg: str, cor: str = TXT_PRI):
        self._log_text.configure(state="normal")
        self._log_text.insert("end", msg + "\n")
        # Cor via tag simulada com inserção colorida não é suportada nativamente
        # no CTkTextbox — usamos texto simples com prefixo visual
        self._log_text.configure(state="disabled")
        self._log_text.see("end")

    def _limpar_log(self):
        self._log_text.configure(state="normal")
        self._log_text.delete("1.0", "end")
        self._log_text.configure(state="disabled")

    # ── Encerramento ──────────────────────────────────────────────────────────

    def _sair(self):
        self._salvar_preferencias()
        self.destroy()



# ════════════════════════════════════════════════════════════════════════════
# Widgets helpers
# ════════════════════════════════════════════════════════════════════════════

def _sep(parent, col):
    ctk.CTkFrame(parent, fg_color=BORDER, width=1, corner_radius=0
                 ).grid(row=0, column=col, sticky="ns")

def _section(parent, text):
    ctk.CTkLabel(parent, text=text, font=FONT_LABEL,
                 text_color=TXT_DIM, anchor="w"
                 ).pack(fill="x", padx=14, pady=(14,4))

def _chk_label(parent) -> ctk.CTkLabel:
    lbl = ctk.CTkLabel(parent, text="...", font=FONT_SMALL,
                       text_color=TXT_DIM, anchor="w")
    lbl.pack(fill="x", padx=22, pady=1)
    return lbl

def _set_chk(lbl: ctk.CTkLabel, ok: bool, txt_ok: str, txt_err: str):
    lbl.configure(
        text=("✔  " if ok else "✖  ") + (txt_ok if ok else txt_err),
        text_color=SUCCESS if ok else (WARN if not ok else TXT_DIM))


# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    app = DDMApp()
    app.mainloop()
