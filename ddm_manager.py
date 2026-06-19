"""
DDM Manager - Metalfrio Solutions
v4.0 — Azul corporativo moderno
"""

# ─── DPI Awareness ───────────────────────────────────────────────────────────
import ctypes, sys
try:
    ctypes.windll.shcore.SetProcessDpiAwareness(2)
except Exception:
    try:
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass

import os, copy, json, shutil, hashlib, datetime, subprocess, tkinter as tk
from tkinter import messagebox, filedialog, simpledialog
from pathlib import Path

# ─── Paleta corporativa ───────────────────────────────────────────────────────
BG_DARK   = "#0A0F1E"   # fundo principal quase preto
BG_CARD   = "#0F1729"   # cards / painéis
BG_INPUT  = "#151E35"   # campos de entrada
BG_ROW    = "#111827"   # linhas da lista
BG_HDR    = "#060D1A"   # header escuro
ACCENT    = "#00D4FF"   # ciano elétrico (destaque principal)
ACCENT2   = "#0EA5E9"   # azul médio
ACCENT3   = "#1CBABE"   # verde-azul Metalfrio
SUCCESS   = "#10B981"   # verde sucesso
WARN_C    = "#F59E0B"   # âmbar alerta
ERR_C     = "#EF4444"   # vermelho erro
GOLD_C    = "#FBBF24"   # dourado admin
TXT_PRI   = "#F1F5F9"   # texto principal
TXT_SEC   = "#94A3B8"   # texto secundário
TXT_DIM   = "#475569"   # texto esmaecido
BORDER    = "#1E293B"   # bordas sutis
SEL_ROW   = "#0C2545"   # linha selecionada/hoje
CHK_BG    = "#0D1F3C"   # fundo checkbox

# Fontes
F_HDR_TITLE = ("Segoe UI", 22, "bold")
F_HDR_SUB   = ("Segoe UI", 10)
F_SECTION   = ("Segoe UI",  8, "bold")
F_LABEL     = ("Segoe UI",  9, "bold")
F_INPUT     = ("Segoe UI", 10)
F_BTN_LG    = ("Segoe UI", 12, "bold")
F_BTN_SM    = ("Segoe UI",  9, "bold")
F_ROW       = ("Segoe UI", 10)
F_ROW_SM    = ("Segoe UI",  8)
F_LOG       = ("Cascadia Code", 9) if sys.platform == "win32" else ("Consolas", 9)
F_BADGE     = ("Segoe UI",  8, "bold")

# Mapeamentos
DIA_MAP  = {0:"2", 1:"3", 2:"4", 3:"5", 4:"6", 5:"7", 6:"7"}
DIAS_PT  = ["Segunda","Terça","Quarta","Quinta","Sexta","Sábado","Domingo"]
DIAS_ABREV = ["SEG","TER","QUA","QUI","SEX","SÁB","DOM"]
MES_MAP  = {
    1:"1 - JANEIRO",  2:"2 - FEVEREIRO", 3:"3 - MARÇO",
    4:"4 - ABRIL",    5:"5 - MAIO",      6:"6 - JUNHO",
    7:"7 - JULHO",    8:"8 - AGOSTO",    9:"9 - SETEMBRO",
    10:"10 - OUTUBRO",11:"11 - NOVEMBRO",12:"12 - DEZEMBRO",
}

PARTICIPANTES = [
    ("10347","ADRIANO GARCIA"),
    ("63363","AGATHA LAÍS SALGUEIRO MAROTTI"),
    ("63051","AMANDA FLORÊNCIO"),
    ("63175","ANA EDUARDA SANO SILVA"),
    ("10934","ANDRÉ LUIZ HIGA FREITAS"),
    ("63538","AUGUSTO GRAÇA SILVESTRIN"),
    ("63231","BARBARA RAMIRES IANHES"),
    ("62334","CLEYTON FARIA ASSIS"),
    ("62397","JOÃO VITOR RODRIGUES SOUZA"),
    ("184",  "JOEL BORGES DE CAMPOS JUNIOR"),
    ("61510","KAROLINE PEDROSO SANTOS"),
    ("10437","LETICIA FERNANDES GONÇALVES"),
    ("63186","LUDIMILA SANTOS SOUZA"),
    ("63516","LUIS GUILHERME DUENHAS SILVA"),
    ("63391","RENAN DA SILVA MANTOVANI"),
    ("63999","RODRIGO ARIAS SILVA"),
    ("11461","VANESSA DIAS DA SILVA"),
    ("9018", "WAGNER JUNIOR ALMEIDA"),
]

DEFAULTS = {
    "raiz":        r"Q:\Transferencia\DDM 2026",
    "setor":       "Engenharia Industrial",
    "subarea":     "Engenharia de Produto",
    "turno":       "ADM",
    "facilitador": "Leitura Individual",
}

# ─── Auth ─────────────────────────────────────────────────────────────────────
USERS_FILE = Path(sys.executable if getattr(sys,"frozen",False)
                  else __file__).parent / "ddm_users.json"

def _hash(pw): return hashlib.sha256(pw.encode()).hexdigest()

def _load_users():
    if USERS_FILE.exists():
        try: return json.loads(USERS_FILE.read_text(encoding="utf-8"))
        except: pass
    users = {"184":{"nome":"JOEL BORGES DE CAMPOS JUNIOR",
                    "pw":_hash("metalfrio"),"admin":True}}
    _save_users(users); return users

def _save_users(u):
    USERS_FILE.write_text(json.dumps(u,ensure_ascii=False,indent=2),encoding="utf-8")

# ════════════════════════════════════════════════════════════════════════════
# Widgets reutilizáveis
# ════════════════════════════════════════════════════════════════════════════

def make_input(parent, var, show=None, width=None):
    """Campo de entrada estilizado."""
    kw = dict(textvariable=var, font=F_INPUT,
              bg=BG_INPUT, fg=TXT_PRI, insertbackground=ACCENT,
              relief="flat", bd=0, highlightthickness=1,
              highlightbackground=BORDER, highlightcolor=ACCENT)
    if show: kw["show"] = show
    if width: kw["width"] = width
    e = tk.Entry(parent, **kw)
    e.bind("<FocusIn>",  lambda _: e.config(highlightbackground=ACCENT))
    e.bind("<FocusOut>", lambda _: e.config(highlightbackground=BORDER))
    return e

def section_label(parent, text):
    """Label de seção com linha decorativa."""
    f = tk.Frame(parent, bg=BG_CARD)
    f.pack(fill="x", pady=(12,4))
    tk.Label(f, text=text.upper(), font=F_SECTION,
             bg=BG_CARD, fg=ACCENT).pack(side="left")
    tk.Frame(f, bg=BORDER, height=1).pack(side="left", fill="x",
                                           expand=True, padx=(8,0), pady=6)
    return f

def accent_btn(parent, text, cmd, primary=False, danger=False, **kw):
    bg = ACCENT if primary else (ERR_C if danger else BG_INPUT)
    fg = BG_DARK if primary else (TXT_PRI if danger else ACCENT)
    b = tk.Button(parent, text=text, command=cmd, font=F_BTN_SM,
                  bg=bg, fg=fg, relief="flat", cursor="hand2",
                  activebackground=ACCENT2, activeforeground=BG_DARK,
                  bd=0, **kw)
    return b

# ════════════════════════════════════════════════════════════════════════════
# Lógica de negócio
# ════════════════════════════════════════════════════════════════════════════

def semana_atual(data):
    seg = data - datetime.timedelta(days=data.weekday())
    return seg, seg + datetime.timedelta(days=4)

def encontrar_ddm_pasta(pasta_dia_path, data):
    nome_mes = MES_MAP[data.month]
    pasta_mes = None
    for item in pasta_dia_path.iterdir():
        if item.is_dir() and (item.name.upper()==nome_mes.upper()
                              or item.name.startswith(str(data.month)+" -")):
            pasta_mes = item; break
    if not pasta_mes:
        return None, f"Pasta '{nome_mes}' não encontrada"
    seg, sex = semana_atual(data)
    alvo, melhor = None, None
    for f in sorted(pasta_mes.glob("*.docx")):
        try:
            dd,mm = int(f.name[:2]),int(f.name[3:5])
            d_arq = datetime.date(data.year,mm,dd)
            if seg <= d_arq <= sex:
                delta = abs((d_arq-data).days)
                if melhor is None or delta < melhor:
                    alvo, melhor = f, delta
        except: continue
    return (alvo,"") if alvo else (None, f"Sem DDM {seg:%d/%m}–{sex:%d/%m}")

PASTA_DIA = {"2":(0,"Segunda-feira"),"3":(1,"Terça-feira"),"4":(2,"Quarta-feira"),
             "5":(3,"Quinta-feira"),"6":(4,"Sexta-feira"),"7":(5,"Sábado")}

def varrer_semana(raiz, data):
    raiz_p = Path(raiz)
    if not raiz_p.exists(): return []
    resultado = []
    for pasta in sorted(raiz_p.iterdir()):
        if not pasta.is_dir(): continue
        partes = pasta.name.split(" - ",1)
        if len(partes)<2: continue
        num = partes[0].strip()
        if num not in PASTA_DIA: continue
        weekday_pasta, dia_nome = PASTA_DIA[num]
        docx, erro = encontrar_ddm_pasta(pasta, data)
        tema = extrair_tema(docx.name) if docx else "—"
        data_ddm = extrair_data_ddm(docx.name if docx else "", data.year)
        resultado.append({"num":int(num),"dia_nome":dia_nome,
                          "dia_abrev":DIAS_ABREV[weekday_pasta],
                          "docx_path":docx,"tema":tema,
                          "data_ddm":data_ddm,"erro":erro})
    resultado.sort(key=lambda x: x["num"])
    return resultado

def extrair_tema(nome):
    stem = Path(nome).stem
    p = stem.split(" ",1)
    return p[1].strip() if len(p)==2 else stem

def extrair_data_ddm(nome, ano):
    try:
        dd,mm = int(nome[:2]),int(nome[3:5])
        return f"{dd:02d}/{mm:02d}/{ano}"
    except: return ""

def _cells_unicas(linha):
    seen,res = set(),[]
    for c in linha.cells:
        cid = id(c._tc)
        if cid not in seen:
            seen.add(cid); res.append(c)
    return res

def _run_apos_label(para, valor):
    from docx.oxml.ns import qn
    from lxml import etree
    runs = para.runs
    if not runs: return
    for r in runs[1:]: r._element.getparent().remove(r._element)
    novo = copy.deepcopy(runs[0]._element)
    t = novo.find(qn("w:t"))
    if t is None: t = etree.SubElement(novo, qn("w:t"))
    t.text = " "+valor
    t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    para._element.append(novo)

def _escrever_celula(cell, valor, alinhar_h="center"):
    from docx.oxml.ns import qn
    from lxml import etree
    tcp = cell._tc.find(qn("w:tcPr"))
    if tcp is None: tcp = etree.SubElement(cell._tc, qn("w:tcPr"))
    val_el = tcp.find(qn("w:vAlign"))
    if val_el is None: val_el = etree.SubElement(tcp, qn("w:vAlign"))
    val_el.set(qn("w:val"),"center")
    para = cell.paragraphs[0] if cell.paragraphs else None
    if not para: return
    for r in para.runs: r._element.getparent().remove(r._element)
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.SubElement(para._element, qn("w:pPr"))
        para._element.insert(0, ppr)
    jc = ppr.find(qn("w:jc"))
    if jc is None: jc = etree.SubElement(ppr, qn("w:jc"))
    jc.set(qn("w:val"),"center" if alinhar_h=="center" else "left")
    r_el = etree.SubElement(para._element, qn("w:r"))
    rpr  = etree.SubElement(r_el, qn("w:rPr"))
    fnts = etree.SubElement(rpr, qn("w:rFonts"))
    for a,v in [("w:ascii","Arial"),("w:hAnsi","Arial"),
                ("w:eastAsia","Batang"),("w:cs","Arial")]:
        fnts.set(qn(a),v)
    for tag in ("w:sz","w:szCs"):
        e = etree.SubElement(rpr, qn(tag)); e.set(qn("w:val"),"18")
    t = etree.SubElement(r_el, qn("w:t"))
    t.text = valor
    t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")

def _data_preenchida(doc):
    for linha in doc.tables[0].rows:
        for cell in _cells_unicas(linha):
            for para in cell.paragraphs:
                txt = para.text.strip()
                if txt.startswith("DATA:") and txt[5:].strip():
                    return True
    return False

def processar_ddm(src, dst, campos):
    from docx import Document
    LABELS = {"SETOR/LINHA:":campos["setor"],"TURNO:":campos["turno"],
              "SUBÁREA:":campos["subarea"],"FACILITADOR:":campos["facilitador"]}
    doc = Document(str(src))
    data_no_doc = _data_preenchida(doc)
    data_extra = {} if data_no_doc else {"DATA:":campos["data"]}
    tabela = doc.tables[0]; part_idx = 0
    info = {"cabecalho":[],"nomes":0,"data_preenchida": not data_no_doc}
    for linha in tabela.rows:
        cells_u = _cells_unicas(linha)
        for cell in cells_u:
            for para in cell.paragraphs:
                txt = para.text.strip()
                for label,valor in {**LABELS,**data_extra}.items():
                    if txt.startswith(label) and txt[len(label):].strip()=="":
                        _run_apos_label(para, valor)
                        info["cabecalho"].append(label)
        if not cells_u: continue
        try: int(cells_u[0].text.strip())
        except ValueError: continue
        if part_idx < len(PARTICIPANTES) and len(cells_u)>=3:
            re_v, nome_v = PARTICIPANTES[part_idx]
            _escrever_celula(cells_u[1], re_v,   alinhar_h="center")
            _escrever_celula(cells_u[2], nome_v, alinhar_h="left")
            info["nomes"] += 1
        part_idx += 1
    doc.save(str(dst))
    return info

def converter_pdfs_word(caminhos):
    try:
        import win32com.client
    except ImportError:
        raise RuntimeError("pywin32 não encontrado. Recompile com 'pywin32'.")
    resultados = []
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        for caminho_docx in caminhos:
            caminho_docx = Path(caminho_docx).resolve()
            caminho_pdf  = caminho_docx.with_suffix(".pdf")
            doc = None
            try:
                doc = word.Documents.Open(str(caminho_docx),
                                          ReadOnly=True, AddToRecentFiles=False)
                doc.SaveAs(str(caminho_pdf), FileFormat=17)
                resultados.append((caminho_pdf,""))
            except Exception as e:
                resultados.append((None, str(e)))
            finally:
                if doc:
                    try: doc.Close(False)
                    except: pass
    finally:
        if word:
            try: word.Quit()
            except: pass
    return resultados

# ════════════════════════════════════════════════════════════════════════════
# Login
# ════════════════════════════════════════════════════════════════════════════

class LoginWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("DDM Manager — Metalfrio")
        self.configure(bg=BG_DARK)
        self.resizable(True, True)
        self.minsize(400, 460)
        self._users = _load_users()
        self._logged_user = None
        self._build()
        self._center(440, 500)

    def _center(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build(self):
        # ── Header ──
        hdr = tk.Frame(self, bg=BG_HDR, height=110)
        hdr.pack(fill="x"); hdr.pack_propagate(False)

        # Linha decorativa ciano no topo
        tk.Frame(hdr, bg=ACCENT, height=3).pack(fill="x")

        inner = tk.Frame(hdr, bg=BG_HDR)
        inner.pack(fill="both", expand=True, padx=28, pady=12)

        row = tk.Frame(inner, bg=BG_HDR)
        row.pack(fill="x", pady=(4,0))
        tk.Label(row, text="❄", font=("Segoe UI", 28),
                 bg=BG_HDR, fg=ACCENT).pack(side="left", padx=(0,10))
        col = tk.Frame(row, bg=BG_HDR)
        col.pack(side="left")
        tk.Label(col, text="DDM Manager", font=("Segoe UI", 20, "bold"),
                 bg=BG_HDR, fg=TXT_PRI).pack(anchor="w")
        tk.Label(col, text="Metalfrio Solutions  ·  Engenharia de Produto",
                 font=F_HDR_SUB, bg=BG_HDR, fg=TXT_SEC).pack(anchor="w")

        # ── Formulário ──
        frm = tk.Frame(self, bg=BG_DARK, padx=32)
        frm.pack(fill="both", expand=True, pady=(24,0))

        # RE
        tk.Label(frm, text="REGISTRO DE EMPREGADO (RE)", font=F_SECTION,
                 bg=BG_DARK, fg=TXT_SEC, anchor="w").pack(fill="x", pady=(0,4))
        self._v_re = tk.StringVar()
        e_re = make_input(frm, self._v_re)
        e_re.pack(fill="x", ipady=9, pady=(0,16))
        e_re.focus()

        # Senha
        tk.Label(frm, text="SENHA", font=F_SECTION,
                 bg=BG_DARK, fg=TXT_SEC, anchor="w").pack(fill="x", pady=(0,4))
        self._v_pw = tk.StringVar()
        e_pw = make_input(frm, self._v_pw, show="●")
        e_pw.pack(fill="x", ipady=9, pady=(0,6))
        e_pw.bind("<Return>", lambda _: self._login())

        self._lbl_err = tk.Label(frm, text="", font=("Segoe UI",9),
                                 bg=BG_DARK, fg=ERR_C)
        self._lbl_err.pack(anchor="w", pady=(0,16))

        # Botão
        btn = tk.Button(frm, text="Entrar  →", font=("Segoe UI",12,"bold"),
                        bg=ACCENT, fg=BG_DARK, relief="flat", cursor="hand2",
                        activebackground=ACCENT2, activeforeground=BG_DARK,
                        command=self._login)
        btn.pack(fill="x", ipady=12)

        # Rodapé
        tk.Label(self, text="© Metalfrio Solutions S.A.",
                 font=("Segoe UI",8), bg=BG_DARK, fg=TXT_DIM
                 ).pack(side="bottom", pady=12)

    def _login(self):
        re = self._v_re.get().strip()
        pw = self._v_pw.get()
        user = self._users.get(re)
        if not user or user["pw"] != _hash(pw):
            self._lbl_err.config(text="  ✕  RE ou senha incorretos.")
            self._v_pw.set(""); return
        self._logged_user = {"re":re, **user}
        self.destroy()

    def get_user(self): return self._logged_user

# ════════════════════════════════════════════════════════════════════════════
# Admin Panel
# ════════════════════════════════════════════════════════════════════════════

class AdminPanel(tk.Toplevel):
    def __init__(self, parent, users):
        super().__init__(parent)
        self.title("Gestão de Usuários")
        self.configure(bg=BG_DARK)
        self.resizable(True, True)
        self.minsize(580, 480)
        self._users = users
        self._build()
        self._center(640, 560)
        self.grab_set()

    def _center(self, w, h):
        self.update_idletasks()
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build(self):
        # Header
        hdr = tk.Frame(self, bg=BG_HDR)
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=GOLD_C, height=3).pack(fill="x")
        inner = tk.Frame(hdr, bg=BG_HDR, padx=20, pady=14)
        inner.pack(fill="x")
        tk.Label(inner, text="⚙  Gestão de Usuários", font=("Segoe UI",14,"bold"),
                 bg=BG_HDR, fg=TXT_PRI).pack(side="left")
        tk.Label(inner, text="Administrador: Joel Borges", font=F_HDR_SUB,
                 bg=BG_HDR, fg=TXT_SEC).pack(side="right")

        # Lista
        lista_frm = tk.Frame(self, bg=BG_DARK, padx=20, pady=12)
        lista_frm.pack(fill="both", expand=True)

        tk.Label(lista_frm, text="USUÁRIOS CADASTRADOS", font=F_SECTION,
                 bg=BG_DARK, fg=TXT_SEC).pack(anchor="w", pady=(0,6))

        # Cabeçalho tabela
        hdr_row = tk.Frame(lista_frm, bg=BG_CARD, padx=8, pady=6)
        hdr_row.pack(fill="x")
        for txt, w in [("RE",8),("Nome",32),("Perfil",8),("Ações",10)]:
            tk.Label(hdr_row, text=txt, font=F_SECTION, bg=BG_CARD,
                     fg=TXT_DIM, width=w, anchor="w").pack(side="left", padx=4)

        self._lista_frm = tk.Frame(lista_frm, bg=BG_DARK)
        self._lista_frm.pack(fill="both", expand=True, pady=(2,0))
        self._refresh_lista()

        # Separador
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x", padx=20)

        # Formulário novo usuário
        novo_frm = tk.Frame(self, bg=BG_DARK, padx=20, pady=14)
        novo_frm.pack(fill="x")
        tk.Label(novo_frm, text="ADICIONAR USUÁRIO", font=F_SECTION,
                 bg=BG_DARK, fg=TXT_SEC).grid(row=0,column=0,columnspan=5,
                                               sticky="w",pady=(0,8))

        self._v_new_re   = tk.StringVar()
        self._v_new_nome = tk.StringVar()
        self._v_new_pw   = tk.StringVar()

        for col,(lbl,var,w,show) in enumerate([
            ("RE", self._v_new_re, 10, None),
            ("Nome completo", self._v_new_nome, 28, None),
            ("Senha", self._v_new_pw, 14, "●"),
        ]):
            tk.Label(novo_frm, text=lbl, font=F_SECTION, bg=BG_DARK,
                     fg=TXT_SEC).grid(row=1,column=col,sticky="w",padx=(0,8))
            e = make_input(novo_frm, var, show=show, width=w)
            e.grid(row=2,column=col,padx=(0,8),ipady=6,sticky="ew")

        tk.Button(novo_frm, text="＋  Adicionar", font=F_BTN_SM,
                  bg=ACCENT, fg=BG_DARK, relief="flat", cursor="hand2",
                  padx=10, command=self._adicionar
                  ).grid(row=2,column=3,padx=(4,0),ipady=6)

        self._lbl_msg = tk.Label(novo_frm, text="", font=("Segoe UI",9),
                                 bg=BG_DARK, fg=SUCCESS)
        self._lbl_msg.grid(row=3,column=0,columnspan=4,sticky="w",pady=(6,0))

        # Rodapé
        rod = tk.Frame(self, bg=BG_HDR, height=48)
        rod.pack(fill="x", side="bottom"); rod.pack_propagate(False)
        tk.Frame(rod, bg=BORDER, height=1).pack(fill="x")
        tk.Button(rod, text="Fechar", font=F_BTN_SM,
                  bg=BG_HDR, fg=TXT_SEC, relief="flat", cursor="hand2",
                  command=self.destroy).pack(side="right", padx=16, pady=10)

    def _refresh_lista(self):
        for w in self._lista_frm.winfo_children(): w.destroy()
        for re, info in self._users.items():
            row = tk.Frame(self._lista_frm, bg=BG_ROW, pady=1)
            row.pack(fill="x", pady=1)
            inner = tk.Frame(row, bg=BG_ROW, padx=8, pady=6)
            inner.pack(fill="x")
            tk.Label(inner, text=re, font=F_INPUT, bg=BG_ROW,
                     fg=TXT_PRI, width=8, anchor="w").pack(side="left", padx=4)
            tk.Label(inner, text=info["nome"][:36], font=F_INPUT, bg=BG_ROW,
                     fg=TXT_PRI, width=32, anchor="w").pack(side="left")
            badge_txt = "ADMIN" if info.get("admin") else "usuário"
            badge_fg  = GOLD_C  if info.get("admin") else TXT_DIM
            tk.Label(inner, text=badge_txt, font=F_BADGE, bg=BG_ROW,
                     fg=badge_fg, width=7).pack(side="left")
            # Ações
            acao_frm = tk.Frame(inner, bg=BG_ROW)
            acao_frm.pack(side="right")
            tk.Button(acao_frm, text="🔑 Senha", font=F_BADGE,
                      bg=BG_CARD, fg=WARN_C, relief="flat", cursor="hand2",
                      padx=6, pady=2,
                      command=lambda r=re: self._redefinir_senha(r)
                      ).pack(side="left", padx=2)
            if not info.get("admin"):
                tk.Button(acao_frm, text="✕ Remover", font=F_BADGE,
                          bg=BG_CARD, fg=ERR_C, relief="flat", cursor="hand2",
                          padx=6, pady=2,
                          command=lambda r=re: self._remover(r)
                          ).pack(side="left", padx=2)

    def _adicionar(self):
        re   = self._v_new_re.get().strip()
        nome = self._v_new_nome.get().strip().upper()
        pw   = self._v_new_pw.get()
        if not re or not nome or not pw:
            self._msg("Preencha RE, nome e senha.", ERR_C); return
        if re in self._users:
            self._msg(f"RE {re} já existe.", ERR_C); return
        self._users[re] = {"nome":nome,"pw":_hash(pw),"admin":False}
        _save_users(self._users)
        self._v_new_re.set(""); self._v_new_nome.set(""); self._v_new_pw.set("")
        self._refresh_lista()
        self._msg(f"Usuário {re} adicionado com sucesso.", SUCCESS)

    def _remover(self, re):
        if not messagebox.askyesno("Confirmar",
                f"Remover {self._users[re]['nome']}?", parent=self): return
        del self._users[re]; _save_users(self._users)
        self._refresh_lista(); self._msg("Usuário removido.", WARN_C)

    def _redefinir_senha(self, re):
        nova = simpledialog.askstring("Redefinir senha",
               f"Nova senha para {self._users[re]['nome']}:",
               show="●", parent=self)
        if not nova: return
        self._users[re]["pw"] = _hash(nova)
        _save_users(self._users)
        self._msg("Senha redefinida com sucesso.", SUCCESS)

    def _msg(self, txt, cor=SUCCESS):
        self._lbl_msg.config(text=txt, fg=cor)

# ════════════════════════════════════════════════════════════════════════════
# Janela Principal
# ════════════════════════════════════════════════════════════════════════════

class MainWindow(tk.Tk):
    def __init__(self, user, users):
        super().__init__()
        self._user  = user
        self._users = users
        self.title("DDM Manager — Metalfrio")
        self.configure(bg=BG_DARK)
        self.resizable(True, True)
        self.minsize(960, 640)
        self._semana  = []
        self._checks  = []
        self._v_saida = tk.StringVar(value="")
        self._build()
        self.after(100, self._varrer)

    def _center(self, w, h):
        self.geometry(f"{w}x{h}")
        self.update_idletasks()
        x = (self.winfo_screenwidth()-w)//2
        y = (self.winfo_screenheight()-h)//2
        self.geometry(f"{w}x{h}+{x}+{y}")

    # ── Build ──────────────────────────────────────────────────────────────
    def _build(self):
        self._center(1100, 720)

        # ══ 1. HEADER ════════════════════════════════════════════════════════
        hdr = tk.Frame(self, bg=BG_HDR)
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=ACCENT, height=3).pack(fill="x")

        hdr_inner = tk.Frame(hdr, bg=BG_HDR, padx=24, pady=14)
        hdr_inner.pack(fill="x")

        # Logo + título
        logo_frm = tk.Frame(hdr_inner, bg=BG_HDR)
        logo_frm.pack(side="left")
        tk.Label(logo_frm, text="❄", font=("Segoe UI", 26),
                 bg=BG_HDR, fg=ACCENT).pack(side="left", padx=(0,12))
        title_frm = tk.Frame(logo_frm, bg=BG_HDR)
        title_frm.pack(side="left")
        tk.Label(title_frm, text="DDM Manager",
                 font=("Segoe UI", 20, "bold"),
                 bg=BG_HDR, fg=TXT_PRI).pack(anchor="w")
        tk.Label(title_frm, text="Metalfrio Solutions  ·  Engenharia de Produto",
                 font=("Segoe UI", 9), bg=BG_HDR, fg=TXT_SEC).pack(anchor="w")

        # Direita do header
        right_frm = tk.Frame(hdr_inner, bg=BG_HDR)
        right_frm.pack(side="right")

        # Badge usuário
        nome_curto = " ".join(self._user["nome"].split()[:2]).title()
        admin_tag  = "  ★  Admin" if self._user.get("admin") else ""
        user_badge = tk.Frame(right_frm, bg=BG_CARD,
                              highlightthickness=1, highlightbackground=BORDER)
        user_badge.pack(side="right", padx=(8,0))
        tk.Label(user_badge, text=f"  👤  {nome_curto}{admin_tag}  ",
                 font=("Segoe UI",9,"bold"), bg=BG_CARD, fg=TXT_PRI,
                 pady=4).pack()

        self._lbl_clock = tk.Label(right_frm, text="",
                                   font=("Segoe UI",9), bg=BG_HDR, fg=TXT_SEC)
        self._lbl_clock.pack(side="right", padx=12)
        self._tick()

        if self._user.get("admin"):
            tk.Button(right_frm, text="⚙  Usuários", font=F_BTN_SM,
                      bg=GOLD_C, fg=BG_DARK, relief="flat", cursor="hand2",
                      padx=12, pady=5,
                      command=self._abrir_admin).pack(side="right", padx=8)

        # ══ 2. RODAPÉ ════════════════════════════════════════════════════════
        rod = tk.Frame(self, bg=BG_HDR, height=68)
        rod.pack(fill="x", side="bottom"); rod.pack_propagate(False)
        tk.Frame(rod, bg=BORDER, height=1).pack(fill="x")

        btn_frm = tk.Frame(rod, bg=BG_HDR)
        btn_frm.pack(fill="both", expand=True, padx=16, pady=10)

        tk.Button(btn_frm, text="📄   Gerar DOCX selecionados",
                  font=F_BTN_LG, bg=BG_INPUT, fg=ACCENT,
                  relief="flat", cursor="hand2", padx=20,
                  activebackground=ACCENT, activeforeground=BG_DARK,
                  command=self._gerar_docx_sel
                  ).pack(side="left", ipady=6)

        tk.Button(btn_frm, text="📋   Combinar DOCX",
                  font=F_BTN_LG, bg=BG_INPUT, fg=ACCENT3,
                  relief="flat", cursor="hand2", padx=20,
                  activebackground=ACCENT3, activeforeground=BG_DARK,
                  command=self._gerar_docx_combinado
                  ).pack(side="left", padx=(8,0), ipady=6)

        tk.Button(btn_frm, text="📑   Gerar PDF + Abrir selecionados",
                  font=F_BTN_LG, bg=ACCENT, fg=BG_DARK,
                  relief="flat", cursor="hand2", padx=20,
                  activebackground=ACCENT2, activeforeground=BG_DARK,
                  command=self._gerar_pdf_sel
                  ).pack(side="left", padx=(8,0), ipady=6)

        tk.Button(btn_frm, text="✖   Sair",
                  font=F_BTN_LG, bg=BG_HDR, fg=ERR_C,
                  relief="flat", cursor="hand2", padx=20,
                  command=self.destroy
                  ).pack(side="right", ipady=6)

        # ══ 3. BARRA PASTA RAIZ ══════════════════════════════════════════════
        paths_bar = tk.Frame(self, bg=BG_CARD, padx=16, pady=8)
        paths_bar.pack(fill="x")
        tk.Frame(paths_bar, bg=BORDER, height=1).pack(fill="x", pady=(0,8))

        # Linha raiz
        raiz_row = tk.Frame(paths_bar, bg=BG_CARD)
        raiz_row.pack(fill="x", pady=(0,4))
        tk.Label(raiz_row, text="📁", font=("Segoe UI",10),
                 bg=BG_CARD, fg=ACCENT).pack(side="left", padx=(0,6))
        tk.Label(raiz_row, text="PASTA RAIZ DDM 2026", font=F_SECTION,
                 bg=BG_CARD, fg=TXT_SEC).pack(side="left", padx=(0,10))

        self._vars: dict[str, tk.StringVar] = {}
        v_raiz = tk.StringVar(value=DEFAULTS.get("raiz",""))
        self._vars["raiz"] = v_raiz
        e_raiz = make_input(raiz_row, v_raiz)
        e_raiz.pack(side="left", fill="x", expand=True, ipady=5, ipadx=6)

        for txt, cmd, primary in [
            ("  …  ", self._browse, True),
            ("🔄  Atualizar", self._varrer, False),
        ]:
            b = accent_btn(raiz_row, txt, cmd, primary=primary)
            b.pack(side="left", padx=(6,0), ipady=5, ipadx=8)

        # Linha saída
        saida_row = tk.Frame(paths_bar, bg=BG_CARD)
        saida_row.pack(fill="x")
        tk.Label(saida_row, text="💾", font=("Segoe UI",10),
                 bg=BG_CARD, fg=TXT_SEC).pack(side="left", padx=(0,6))
        tk.Label(saida_row, text="SALVAR EM", font=F_SECTION,
                 bg=BG_CARD, fg=TXT_SEC).pack(side="left", padx=(0,10))
        self._lbl_saida = tk.Label(saida_row, textvariable=self._v_saida,
                                   font=F_INPUT, bg=BG_CARD, fg=TXT_PRI, anchor="w")
        self._lbl_saida.pack(side="left", fill="x", expand=True)
        tk.Label(saida_row, text="(padrão: mesma pasta do DDM)",
                 font=("Segoe UI",8), bg=BG_CARD, fg=TXT_DIM
                 ).pack(side="left", padx=8)
        accent_btn(saida_row, "Alterar", self._browse_saida
                   ).pack(side="left", padx=(0,4), ipady=4, ipadx=8)
        accent_btn(saida_row, "✕", lambda: self._v_saida.set(""),
                   danger=True).pack(side="left", ipady=4, ipadx=6)

        tk.Frame(paths_bar, bg=BORDER, height=1).pack(fill="x", pady=(8,0))

        # ══ 4. CORPO ══════════════════════════════════════════════════════════
        body = tk.Frame(self, bg=BG_DARK)
        body.pack(fill="both", expand=True)

        # ── Coluna esquerda: campos ──
        esq = tk.Frame(body, bg=BG_CARD, width=290)
        esq.pack(side="left", fill="y")
        esq.pack_propagate(False)
        tk.Frame(esq, bg=BORDER, width=1).pack(side="right", fill="y")

        esq_inner = tk.Frame(esq, bg=BG_CARD, padx=16)
        esq_inner.pack(fill="both", expand=True, pady=8)

        section_label(esq_inner, "Cabeçalho DDM")

        for key, label, icon in [
            ("setor",       "Setor / Linha",  "🏭"),
            ("subarea",     "Subárea",        "🔹"),
            ("turno",       "Turno",          "⏰"),
            ("facilitador", "Facilitador",    "👤"),
        ]:
            lbl_row = tk.Frame(esq_inner, bg=BG_CARD)
            lbl_row.pack(fill="x", pady=(8,2))
            tk.Label(lbl_row, text=icon, font=("Segoe UI",9),
                     bg=BG_CARD, fg=TXT_SEC).pack(side="left", padx=(0,4))
            tk.Label(lbl_row, text=label.upper(), font=F_SECTION,
                     bg=BG_CARD, fg=TXT_SEC).pack(side="left")
            v = tk.StringVar(value=DEFAULTS.get(key,""))
            self._vars[key] = v
            e = make_input(esq_inner, v)
            e.pack(fill="x", ipady=7, ipadx=6)

        # Badge participantes
        section_label(esq_inner, "Participantes")
        badge = tk.Frame(esq_inner, bg=BG_INPUT,
                         highlightthickness=1, highlightbackground=BORDER)
        badge.pack(fill="x", pady=4)
        tk.Label(badge, text=f"👥  {len(PARTICIPANTES)} colaboradores carregados",
                 font=("Segoe UI",9), bg=BG_INPUT, fg=SUCCESS,
                 pady=8).pack()

        # ── Coluna central: DDMs da semana ──
        meio = tk.Frame(body, bg=BG_DARK, width=380)
        meio.pack(side="left", fill="y")
        meio.pack_propagate(False)
        tk.Frame(meio, bg=BORDER, width=1).pack(side="right", fill="y")

        meio_inner = tk.Frame(meio, bg=BG_DARK, padx=14)
        meio_inner.pack(fill="both", expand=True, pady=8)

        # Título semana
        seg, sex = semana_atual(datetime.date.today())
        titulo_row = tk.Frame(meio_inner, bg=BG_DARK)
        titulo_row.pack(fill="x", pady=(4,8))
        tk.Label(titulo_row, text="SEMANA", font=F_SECTION,
                 bg=BG_DARK, fg=TXT_DIM).pack(side="left")
        tk.Label(titulo_row,
                 text=f"  {seg:%d/%m} – {sex:%d/%m}",
                 font=("Segoe UI",9,"bold"), bg=BG_DARK, fg=ACCENT
                 ).pack(side="left")

        # Cabeçalho lista
        hdr_list = tk.Frame(meio_inner, bg=BG_CARD, padx=8, pady=5)
        hdr_list.pack(fill="x")
        for txt, w in [("",3),("DIA",14),("TEMA",20),("DATA",10)]:
            tk.Label(hdr_list, text=txt, font=F_SECTION, bg=BG_CARD,
                     fg=TXT_DIM, width=w, anchor="w").pack(side="left", padx=3)

        self._lista_frm = tk.Frame(meio_inner, bg=BG_DARK)
        self._lista_frm.pack(fill="both", expand=True, pady=(2,0))

        # Botões seleção
        sel_row = tk.Frame(meio_inner, bg=BG_DARK)
        sel_row.pack(fill="x", pady=(8,0))
        for txt, val in [("Selecionar tudo", True), ("Limpar seleção", False)]:
            fg = ACCENT if val else TXT_DIM
            tk.Button(sel_row, text=txt, font=F_BTN_SM,
                      bg=BG_CARD, fg=fg, relief="flat", cursor="hand2",
                      padx=10, pady=4,
                      command=lambda v=val: [x.set(v) for x in self._checks]
                      ).pack(side="left", padx=(0,4))

        # ── Coluna direita: log ──
        dir_ = tk.Frame(body, bg=BG_DARK)
        dir_.pack(side="right", fill="both", expand=True)

        log_hdr = tk.Frame(dir_, bg=BG_CARD, padx=16, pady=8)
        log_hdr.pack(fill="x")
        tk.Label(log_hdr, text="LOG DE EXECUÇÃO", font=F_SECTION,
                 bg=BG_CARD, fg=TXT_DIM).pack(side="left")
        tk.Button(log_hdr, text="Limpar", font=F_BADGE,
                  bg=BG_CARD, fg=TXT_DIM, relief="flat", cursor="hand2",
                  command=self._log_clear).pack(side="right")

        self._log_w = tk.Text(dir_, font=F_LOG, bg=BG_DARK, fg=TXT_PRI,
                              relief="flat", state="disabled", wrap="word",
                              padx=12, pady=8,
                              selectbackground=ACCENT, selectforeground=BG_DARK)
        self._log_w.pack(fill="both", expand=True)

    # ── Tick ──────────────────────────────────────────────────────────────
    def _tick(self):
        agora = datetime.datetime.now()
        self._lbl_clock.config(
            text=f"{DIAS_PT[agora.weekday()]}  {agora.strftime('%d/%m/%Y   %H:%M:%S')}"
        )
        self.after(1000, self._tick)

    def _browse(self):
        p = filedialog.askdirectory(title="Pasta raiz DDM 2026")
        if p: self._vars["raiz"].set(p); self._varrer()

    def _browse_saida(self):
        p = filedialog.askdirectory(title="Onde salvar os arquivos gerados")
        if p: self._v_saida.set(p)

    def _abrir_admin(self):
        AdminPanel(self, self._users)

    # ── Varrer semana ─────────────────────────────────────────────────────
    def _varrer(self):
        self._log_clear()
        raiz = self._vars["raiz"].get().strip()
        hoje = datetime.date.today()
        seg, sex = semana_atual(hoje)
        self._log(f"Semana: {seg:%d/%m/%Y} – {sex:%d/%m/%Y}", TXT_SEC)
        self._log(f"Raiz:   {raiz}\n", TXT_SEC)
        self._semana = varrer_semana(raiz, hoje)
        self._rebuild_lista()
        ok  = [d for d in self._semana if d["docx_path"]]
        err = [d for d in self._semana if not d["docx_path"]]
        self._log(f"✔  {len(ok)} DDM(s) encontrado(s)", SUCCESS)
        for d in ok:
            self._log(f"   {d['dia_nome']}: {d['docx_path'].name}", TXT_SEC)
        if err:
            self._log(f"\n⚠  {len(err)} pasta(s) sem DDM:", WARN_C)
            for d in err:
                self._log(f"   {d['dia_nome']}: {d['erro']}", TXT_DIM)

    # ── Lista ─────────────────────────────────────────────────────────────
    def _rebuild_lista(self):
        for w in self._lista_frm.winfo_children(): w.destroy()
        self._checks = []
        hoje = datetime.date.today()

        for d in self._semana:
            tem = d["docx_path"] is not None
            v   = tk.BooleanVar(value=tem)
            self._checks.append(v)

            eh_hoje = (DIA_MAP.get(hoje.weekday()) == str(d["num"]))
            bg = SEL_ROW if eh_hoje else BG_ROW

            row = tk.Frame(self._lista_frm, bg=bg,
                           cursor="hand2" if tem else "")
            row.pack(fill="x", pady=1)

            # Borda esquerda colorida no dia atual
            if eh_hoje:
                tk.Frame(row, bg=ACCENT, width=3).pack(side="left", fill="y")

            inner = tk.Frame(row, bg=bg, padx=6, pady=7)
            inner.pack(side="left", fill="x", expand=True)

            # Checkbox canvas
            cv = tk.Canvas(inner, width=16, height=16, bg=bg,
                           highlightthickness=0,
                           cursor="hand2" if tem else "")
            cv.pack(side="left", padx=(0,8))

            def _draw(canvas, var, enabled, bg_=bg):
                canvas.delete("all")
                if enabled:
                    if var.get():
                        canvas.create_rectangle(1,1,15,15,
                            outline=ACCENT, fill=ACCENT, width=1)
                        canvas.create_line(3,8,7,12, fill=BG_DARK, width=2)
                        canvas.create_line(7,12,13,4, fill=BG_DARK, width=2)
                    else:
                        canvas.create_rectangle(1,1,15,15,
                            outline=TXT_DIM, fill=bg_, width=1)
                else:
                    canvas.create_rectangle(1,1,15,15,
                        outline=TXT_DIM, fill=bg_, width=1, dash=(2,2))

            _draw(cv, v, tem)

            def _toggle(event, var=v, canvas=cv, enabled=tem, bg_=bg):
                if not enabled: return
                var.set(not var.get())
                _draw(canvas, var, enabled, bg_)

            for widget in (row, inner, cv):
                widget.bind("<Button-1>", _toggle)

            # Abreviação dia
            fg_dia = ACCENT if eh_hoje else (TXT_PRI if tem else TXT_DIM)
            lbl_dia = tk.Label(inner, text=d["dia_abrev"],
                               font=("Segoe UI",9,"bold"),
                               bg=bg, fg=fg_dia, width=5, anchor="w")
            lbl_dia.pack(side="left")
            lbl_dia.bind("<Button-1>", _toggle)

            if tem:
                tema_txt = d["tema"][:28]+"…" if len(d["tema"])>28 else d["tema"]
                lbl_t = tk.Label(inner, text=tema_txt, font=F_ROW,
                                 bg=bg, fg=TXT_PRI, anchor="w")
                lbl_t.pack(side="left", fill="x", expand=True)
                lbl_t.bind("<Button-1>", _toggle)
                lbl_d = tk.Label(inner, text=d["data_ddm"],
                                 font=("Segoe UI",8), bg=bg, fg=TXT_DIM)
                lbl_d.pack(side="right", padx=8)
                lbl_d.bind("<Button-1>", _toggle)
            else:
                lbl_na = tk.Label(inner, text="sem arquivo",
                                  font=("Segoe UI",8,"italic"),
                                  bg=bg, fg=TXT_DIM)
                lbl_na.pack(side="left")
                lbl_na.bind("<Button-1>", _toggle)

    # ── Helpers ───────────────────────────────────────────────────────────
    def _campos_para(self, item):
        data_ddm = item["data_ddm"] or datetime.date.today().strftime("%d/%m/%Y")
        return {
            "setor":       self._vars["setor"].get(),
            "subarea":     self._vars["subarea"].get(),
            "turno":       self._vars["turno"].get(),
            "facilitador": self._vars["facilitador"].get(),
            "data":        data_ddm,
        }

    def _pasta_saida(self, fallback):
        v = self._v_saida.get().strip()
        if v:
            p = Path(v); p.mkdir(parents=True, exist_ok=True); return p
        return fallback

    def _selecionados(self):
        sel = [d for d,v in zip(self._semana, self._checks)
               if v.get() and d["docx_path"]]
        if not sel:
            messagebox.showwarning("Nenhum selecionado",
                "Selecione ao menos um DDM da lista.")
        return sel

    # ── Gerar DOCX ────────────────────────────────────────────────────────
    def _gerar_docx_sel(self):
        sel = self._selecionados()
        if not sel: return
        self._log(f"\n─── Gerando {len(sel)} DOCX(s) ───", ACCENT)
        gerados = []
        for d in sel:
            src   = d["docx_path"]
            pasta = self._pasta_saida(src.parent)
            dst   = pasta / (src.stem + "_PREENCHIDO.docx")
            try:
                info = processar_ddm(src, dst, self._campos_para(d))
                self._log(f"✔  {d['dia_nome']}: {dst.name}", SUCCESS)
                self._log(f"   {info['nomes']} participantes · "
                          f"{len(info['cabecalho'])} campos", TXT_SEC)
                gerados.append(dst)
            except Exception as e:
                self._log(f"✖  {d['dia_nome']}: {e}", ERR_C)
        if gerados and messagebox.askyesno("Concluído",
                f"{len(gerados)} DOCX(s) gerado(s).\nAbrir pasta?"):
            os.startfile(str(gerados[0].parent))

    # ── Gerar PDF ─────────────────────────────────────────────────────────
    def _gerar_pdf_sel(self):
        sel = self._selecionados()
        if not sel: return
        self._log(f"\n─── Gerando {len(sel)} PDF(s) ───", ACCENT)

        docxs = []
        for d in sel:
            src   = d["docx_path"]
            pasta = self._pasta_saida(src.parent)
            docx  = pasta / (src.stem + "_PREENCHIDO.docx")
            try:
                info = processar_ddm(src, docx, self._campos_para(d))
                self._log(f"✔  {d['dia_nome']}: {info['nomes']} participantes", SUCCESS)
                docxs.append((d["dia_nome"], docx))
            except Exception as e:
                self._log(f"✖  {d['dia_nome']} (DOCX): {e}", ERR_C)

        if not docxs: return

        self._log(f"\n   Iniciando Word ({len(docxs)} arquivo(s))...", TXT_SEC)
        try:
            resultados = converter_pdfs_word([p for _,p in docxs])
        except Exception as e:
            self._log(f"✖  Word: {e}", ERR_C); return

        pdfs = []
        for (dia_nome,_),(pdf_path,erro) in zip(docxs, resultados):
            if pdf_path:
                self._log(f"✔  {dia_nome}: {pdf_path.name}", SUCCESS)
                pdfs.append(pdf_path)
            else:
                self._log(f"✖  {dia_nome}: {erro}", ERR_C)

        if not pdfs: return

        if len(pdfs) == 1:
            os.startfile(str(pdfs[0]))
        else:
            self._log(f"\n   Mesclando {len(pdfs)} PDFs...", TXT_SEC)
            try:
                pasta_merged = self._pasta_saida(pdfs[0].parent)
                merged = self._mesclar_pdfs(pdfs, pasta_merged)
                self._log(f"✔  PDF combinado: {merged.name}", SUCCESS)
                os.startfile(str(merged))
            except Exception as e:
                self._log(f"⚠  Mesclagem falhou ({e}). Abrindo individualmente.", WARN_C)
                for p in pdfs: os.startfile(str(p))

    def _mesclar_pdfs(self, pdfs, pasta_saida=None):
        try:
            from pypdf import PdfWriter
        except ImportError:
            from PyPDF2 import PdfWriter
        writer = PdfWriter()
        for p in pdfs: writer.append(str(p))
        saida = (pasta_saida or pdfs[0].parent) / "DDM_SEMANA_COMBINADO.pdf"
        with open(saida,"wb") as f: writer.write(f)
        return saida

    def _gerar_docx_combinado(self):
        """Gera um único DOCX com todos os DDMs selecionados, separados por quebra de página."""
        sel = self._selecionados()
        if not sel: return
        if len(sel) == 1:
            messagebox.showinfo("Atenção",
                "Apenas um DDM selecionado.\nUse 'Gerar DOCX selecionados' para um único arquivo.")
            return

        self._log(f"\n─── Combinando {len(sel)} DOCX(s) ───", ACCENT3)

        # 1. Gera cada DOCX preenchido individualmente
        docxs_prontos = []
        for d in sel:
            src   = d["docx_path"]
            pasta = self._pasta_saida(src.parent)
            dst   = pasta / (src.stem + "_PREENCHIDO.docx")
            try:
                info = processar_ddm(src, dst, self._campos_para(d))
                self._log(f"✔  {d['dia_nome']}: {info['nomes']} participantes", SUCCESS)
                docxs_prontos.append((d["dia_nome"], dst))
            except Exception as e:
                self._log(f"✖  {d['dia_nome']}: {e}", ERR_C)

        if len(docxs_prontos) < 2:
            self._log("⚠  Menos de 2 arquivos gerados — combinação cancelada.", WARN_C)
            return

        # 2. Mescla em um único DOCX
        self._log(f"\n   Mesclando em um único DOCX...", TXT_SEC)
        try:
            pasta_out = self._pasta_saida(docxs_prontos[0][1].parent)
            saida = self._mesclar_docxs([p for _,p in docxs_prontos], pasta_out)
            self._log(f"✔  DOCX combinado: {saida.name}", SUCCESS)
            if messagebox.askyesno("Concluído",
                    f"DOCX combinado gerado:\n{saida.name}\n\nAbrir?"):
                os.startfile(str(saida))
        except Exception as e:
            self._log(f"✖  Erro na combinação: {e}", ERR_C)

    def _mesclar_docxs(self, docxs: list, pasta_saida: Path) -> Path:
        """
        Mescla lista de .docx em um único arquivo.
        Insere quebra de página entre cada documento.
        """
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from lxml import etree
        import zipfile, shutil, re as _re

        # Usa o primeiro como base
        base_path = docxs[0]
        saida = pasta_saida / "DDM_SEMANA_COMBINADO.docx"
        shutil.copy(str(base_path), str(saida))

        doc_base = Document(str(saida))

        for docx_path in docxs[1:]:
            # Quebra de página explícita antes de cada novo DDM
            ultimo_para = doc_base.paragraphs[-1] if doc_base.paragraphs else None
            if ultimo_para:
                pPr = ultimo_para._element.get_or_add_pPr()
                pb = OxmlElement("w:pageBreakBefore")
                pb.set(qn("w:val"), "false")
                pPr.append(pb)

            # Adiciona parágrafo com quebra de página
            para_break = doc_base.add_paragraph()
            run = para_break.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run._element.append(br)

            # Copia o conteúdo do próximo doc
            doc_add = Document(str(docx_path))

            # Copia tabelas (o DDM é basicamente uma tabela)
            for tabela in doc_add.tables:
                # Serializa a tabela como XML e insere no doc base
                tbl_xml = copy.deepcopy(tabela._tbl)
                doc_base.element.body.append(tbl_xml)

            # Copia parágrafos soltos (fora de tabelas) se houver
            body_add = doc_add.element.body
            for child in body_add:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag == "p":
                    txt = "".join(c.text or "" for c in child.iter()
                                  if c.tag.endswith("}t"))
                    if txt.strip():
                        doc_base.element.body.append(copy.deepcopy(child))

        doc_base.save(str(saida))
        return saida

    # ── Log ───────────────────────────────────────────────────────────────
    def _log(self, msg, cor=None):
        cor = cor or TXT_PRI
        w = self._log_w
        w.configure(state="normal")
        tag = "c"+cor.replace("#","")
        w.tag_configure(tag, foreground=cor)
        w.insert("end", msg+"\n", tag)
        w.see("end")
        w.configure(state="disabled")

    def _log_clear(self):
        self._log_w.configure(state="normal")
        self._log_w.delete("1.0","end")
        self._log_w.configure(state="disabled")

# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    login = LoginWindow()
    login.mainloop()
    user = login.get_user()
    if not user: sys.exit(0)
    users = _load_users()
    app = MainWindow(user, users)
    app.mainloop()
